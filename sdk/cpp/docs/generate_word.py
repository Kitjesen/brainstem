"""
generate_word.py — 生成 ORIX C++ SDK 开发手册 (.docx)

依赖：pip install python-docx
运行：cd sdk/cpp/docs && python generate_word.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


# ── 样式工具 ──────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def code_block(doc, code_text):
    """等宽代码块，灰色背景"""
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)
    return p


def add_table(doc, headers, rows, header_color="DBEAFE"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_color)
        cell._tc.get_or_add_tcPr().append(shd)
    for r, row_data in enumerate(rows):
        row = table.rows[r + 1]
        for c, val in enumerate(row_data):
            row.cells[c].text = str(val)
    return table


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def warning_box(doc, text):
    """橙色警告框"""
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run("WARNING: " + text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC8, 0x78, 0x00)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF3CD")
    p._p.get_or_add_pPr().append(shd)
    return p


def tip_box(doc, text):
    """蓝绿色提示框"""
    p = doc.add_paragraph()
    p.style = "No Spacing"
    run = p.add_run("TIP: " + text)
    run.font.color.rgb = RGBColor(0x08, 0x91, 0xB2)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E0F7FA")
    p._p.get_or_add_pPr().append(shd)
    return p


# ── 文档构建 ──────────────────────────────────────────────────

def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # ================================================================
    #  封面
    # ================================================================
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ORIX C++ SDK 开发手册")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x08, 0x91, 0xB2)

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("面向高性能嵌入式部署的专业级机器人开发平台")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = brand.add_run("大算机器人")
    r.bold = True
    r.font.size = Pt(18)

    doc.add_paragraph()

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ver.add_run("Brainstem SDK for C++  v1.0.0")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(
        f"日期: {datetime.date.today().strftime('%Y-%m-%d')}"
    ).font.size = Pt(11)

    doc.add_paragraph()

    copy_p = doc.add_paragraph()
    copy_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = copy_p.add_run("Copyright (c) 2026 大算机器人. All rights reserved.")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # ── 版权页 ───────────────────────────────────────────────
    set_heading(doc, "文档信息", 2)
    add_table(doc, ["项目", "说明"], [
        ["文档名称", "ORIX 教育版四足机器狗 C++ SDK 开发手册"],
        ["版本号", "v1.0.0"],
        ["编写日期", "2026年4月"],
        ["适用产品", "ORIX 教育版四足机器狗"],
        ["适用 SDK", "orix (C++)"],
    ])
    doc.add_paragraph()

    set_heading(doc, "版本历史", 2)
    add_table(doc, ["版本", "日期", "作者", "说明"], [
        ["v1.0.0", "2026-04", "大算机器人", "正式发布"],
    ])
    doc.add_paragraph()
    doc.add_paragraph("本文档为大算机器人内部技术文档，未经许可不得外传。")
    doc.add_page_break()

    # ── 前言 ─────────────────────────────────────────────────
    set_heading(doc, "前言", 1)
    doc.add_paragraph(
        "近年来，国家高度重视人工智能与机器人教育的发展。四足机器人是一类典型的"
        "高集成度教学平台，融合了机械工程、电子控制、嵌入式系统、计算机视觉和人工"
        "智能等多个学科的核心内容。"
    )
    doc.add_paragraph(
        "ORIX 教育版四足机器狗是大算机器人面向教育场景推出的专业级开发平台。"
        "我们的设计遵循四个核心理念：学以致用——它不是展会上的观摩样品，而是一台"
        "可以编程、调参、训练自定义策略的真实机器人；降低门槛——几行 C++ 代码即可"
        "完成从连接到行走的全流程；不设天花板——从入门的遥控行走到高级的 RL 策略"
        "训练、SLAM 建图和自主导航，一个平台覆盖完整学习路径；工业级可靠——采用"
        "与工业机器人相同的 FSM 状态机、控制权仲裁和 E-Stop 安全机制。"
    )
    doc.add_paragraph(
        "本手册是 ORIX C++ SDK 的完整开发指南。相较于 Python SDK，C++ SDK 面向"
        "以下场景：高性能嵌入式部署——直接在机器人板端编译运行，消除网络延迟；"
        "ROS2 节点开发——作为 ROS2 包的组成部分；低延迟控制——无 GIL 限制；"
        "与 OpenCV 等 C++ 库无缝协作。手册覆盖从环境安装到高级开发的完整流程，"
        "建议首次使用时按顺序阅读前三章，完成首次连通后再根据兴趣阅读后续章节。"
    )
    doc.add_paragraph()

    # ================================================================
    #  第 1 章：概述
    # ================================================================
    set_heading(doc, "第一章 概述", 1)

    set_heading(doc, "1.1 产品定位", 2)
    doc.add_paragraph(
        "ORIX 教育版四足机器狗是一个集成了运动控制、传感器融合、强化学习策略推理"
        "和安全机制的完整机器人平台。它由机器狗本体（12 DOF 足式四足机器人）、C++ SDK"
        "（orix）和开发文档三部分组成。机器狗内置完整的控制系统，开发者通过 SDK "
        "在开发电脑或机器人本体上编写 C++ 代码控制机器人，无需接触底层通信协议。"
    )
    doc.add_paragraph(
        "C++ SDK 特别适合需要极低延迟或嵌入式部署的场景。与 Python SDK 功能完全"
        "对等，但提供了原生性能、无 GIL 限制和与 OpenCV/ROS2 等 C++ 生态的直接"
        "集成能力。"
    )

    set_heading(doc, "1.2 SDK 能做什么", 2)
    doc.add_paragraph("orix C++ SDK 提供了对机器狗的全方位控制能力，主要包括以下五类接口：")
    bullet(doc, "运动控制 --- 使能/禁用电机、站立、坐下、行走、切换速度模式、播放预设动作")
    bullet(doc, "传感器读取 --- IMU 姿态数据流（50Hz）、关节编码器数据流、电机状态、母线电压")
    bullet(doc, "状态管理 --- 查询 FSM 状态、订阅状态变化事件、切换运动策略")
    bullet(doc, "诊断与维护 --- 读取电机故障码、清除故障、电机标零、健康检查")
    bullet(doc, "摄像头接口 --- USB 摄像头拍照、录像、MJPEG 流媒体服务（需 OpenCV，可选模块）")
    doc.add_paragraph(
        "整个 SDK 遵循简洁设计原则：客户代码完全不需要了解 gRPC、protobuf、"
        "状态机实现细节，所有复杂性隐藏在 orix::OrixClient 类后面。"
    )
    code_block(doc, """\
#include "orix/orix.h"

int main() {
    orix::OrixClient dog("192.168.66.190");
    dog.enable();
    dog.stand_up();
    dog.walk(0.3);   // 前进 30% 速度
    dog.sit_down();
    dog.disable();
    return 0;
}\
""")

    set_heading(doc, "1.3 典型开发流程", 2)
    numbered(doc, "环境准备 --- 开发机安装 C++17 编译器、CMake 3.16+、gRPC 1.50+ 和 Protobuf 3.21+")
    numbered(doc, "安装 SDK --- 通过 vcpkg 安装依赖，或手动 cmake 编译")
    numbered(doc, "验证连接 --- 编译并运行验证程序，确认与机器狗通信正常")
    numbered(doc, "编写控制代码 --- 使用 orix::OrixClient 发送运动指令和读取传感器数据")
    numbered(doc, "调试与迭代 --- 根据需要切换策略、调整参数、查看日志")
    numbered(doc, "高级开发 --- 集成 OpenCV 计算机视觉、开发 ROS2 节点、训练自定义 RL 策略")

    set_heading(doc, "1.4 手册结构", 2)
    doc.add_paragraph("本手册共 11 章，按\"入门 -> 理论 -> 参考 -> 运维\"的顺序组织：")
    add_table(doc, ["章", "主题", "适合谁"], [
        ["1", "概述（本章）", "所有读者"],
        ["2", "快速开始", "第一次使用的开发者"],
        ["3", "理论基础", "希望理解控制原理的学生/研究人员"],
        ["4", "硬件概览", "需要了解硬件参数的工程师"],
        ["5", "坐标系与关节", "进行运动规划的开发者"],
        ["6", "状态机与运动控制", "编写控制逻辑的开发者"],
        ["7", "实用开发指南", "有实际项目需求的 C++ 开发者"],
        ["8", "传感器数据", "进行数据采集和算法研究的开发者"],
        ["9", "API 参考", "查阅接口细节时的所有开发者"],
        ["10", "安全指南", "上手前必读"],
        ["11", "常见问题", "遇到问题时查阅"],
    ])
    doc.add_paragraph(
        "首次使用建议按顺序阅读第 1-2 章（概述 + 快速开始），完成首次连通和基本控制验证。"
        "之后可根据需要跳读。安全指南（第 10 章）在真机操作前必读。"
    )
    doc.add_paragraph()

    # ================================================================
    #  第 2 章：快速开始
    # ================================================================
    set_heading(doc, "第二章 快速开始", 1)

    set_heading(doc, "2.1 环境要求", 2)
    doc.add_paragraph("在开始编译 SDK 之前，请确认您的开发环境满足以下最低要求：")
    add_table(doc, ["项目", "要求"], [
        ["C++ 标准", "C++17 或更高"],
        ["编译器", "GCC 9+ / Clang 10+ / MSVC 2019+"],
        ["构建系统", "CMake 3.16+"],
        ["gRPC", ">= 1.50.0"],
        ["Protobuf", ">= 3.21.0"],
        ["操作系统", "Linux / macOS / Windows"],
        ["网络", "开发机与机器人在同一局域网"],
        ["端口", "TCP 13145（gRPC over HTTP/2）"],
    ])
    warning_box(doc, "不需要安装 ROS、消息中间件或其他额外组件。SDK 通过纯 gRPC 直连机器人。")
    doc.add_paragraph()

    set_heading(doc, "2.2 安装与编译", 2)

    set_heading(doc, "2.2.1 方式一：vcpkg 安装依赖（推荐）", 3)
    code_block(doc, """\
# 安装 vcpkg（如果尚未安装）
git clone https://github.com/microsoft/vcpkg.git
cd vcpkg && ./bootstrap-vcpkg.sh

# 安装 gRPC 和 Protobuf
./vcpkg install grpc protobuf

# 如果需要摄像头功能，同时安装 OpenCV
./vcpkg install opencv4\
""")

    set_heading(doc, "2.2.2 方式二：CMake 编译步骤", 3)
    code_block(doc, """\
cd sdk/cpp
mkdir build && cd build
cmake .. -DCMAKE_BUILD_TYPE=Release
cmake --build . --config Release\
""")
    doc.add_paragraph("如果使用 vcpkg：")
    code_block(doc, """\
cmake .. -DCMAKE_BUILD_TYPE=Release \\
         -DCMAKE_TOOLCHAIN_FILE=/path/to/vcpkg/scripts/buildsystems/vcpkg.cmake
cmake --build . --config Release\
""")

    set_heading(doc, "2.2.3 验证编译", 3)
    code_block(doc, """\
# 编译后可执行文件在 build/examples/ 目录下
./build/examples/01_hello_walk --help\
""")
    doc.add_paragraph()

    set_heading(doc, "2.3 三行代码控制机器人", 2)
    doc.add_paragraph(
        "下面这段代码适合作为\"环境和链路是否通了\"的最小验证，而不是正式的运动脚本。"
    )
    code_block(doc, """\
#include "orix/orix.h"

int main() {
    orix::OrixClient dog("192.168.66.190");
    dog.walk(0.3);  // 30% 速度前进
    return 0;
}\
""")
    warning_box(doc,
        "实际使用中，您需要先 enable() 使能电机并 stand_up() 站起来，才能执行 walk()。"
        "完整的操作流程见下方\"第一次完整操作\"。"
    )
    doc.add_paragraph()

    set_heading(doc, "2.4 首次连接检查清单", 2)
    doc.add_paragraph("在第一次连接机器人之前，请逐项确认：")
    bullet(doc, "机器人已上电，电池电压正常（正常工作电压约 24V，低压保护 18V）")
    bullet(doc, "机器人主控板（S100P）已启动完成（约 30 秒）")
    bullet(doc, "brainstem 服务已运行（systemctl status brainstem）")
    bullet(doc, "开发机与机器人在同一局域网（可以 ping 192.168.66.190）")
    bullet(doc, "遥控器已关闭或置于空闲状态（遥控器优先级高于 SDK）")
    bullet(doc, "机器人放置在平坦地面上，四足触地")
    doc.add_paragraph()

    set_heading(doc, "2.5 网络配置", 2)
    doc.add_paragraph(
        "开发机（PC 或笔记本，IP 为 192.168.66.xxx）与 ORIX 机器人（IP 为 192.168.66.190）"
        "需处于同一局域网内。SDK 使用 TCP 端口 13145 进行 gRPC 通信，无需额外中间件。"
    )
    add_table(doc, ["参数", "默认值", "说明"], [
        ["机器人 IP", "192.168.66.190", "S100P 主控板默认地址"],
        ["gRPC 端口", "13145", "无需修改"],
        ["协议", "gRPC (HTTP/2)", "无加密（局域网内使用）"],
        ["超时", "5.0 秒", "可通过构造函数 timeout 参数调整"],
    ])
    doc.add_paragraph("如果机器人 IP 不同，连接时指定即可：")
    code_block(doc, 'orix::OrixClient dog("10.0.0.100", 13145, 5.0);')
    doc.add_paragraph()

    set_heading(doc, "2.6 验证连接程序", 2)
    doc.add_paragraph("首次连接机器狗时，建议使用以下验证程序逐项检查通信状态：")
    code_block(doc, """\
#include "orix/orix.h"
#include <iostream>

int main() {
    const std::string ROBOT_IP = "192.168.66.190";
    std::cout << "正在连接 ORIX @ " << ROBOT_IP << ":13145 ...\\n";

    try {
        orix::OrixClient dog(ROBOT_IP, 13145, 3.0);

        auto state = dog.get_state();
        std::cout << "[OK] 状态机: " << orix::to_string(state) << "\\n";

        auto voltages = dog.get_voltage();
        if (!voltages.empty()) {
            double avg = 0;
            for (auto v : voltages) avg += v;
            avg /= voltages.size();
            std::cout << "[OK] 平均电压: " << avg << "V ("
                      << voltages.size() << " 个电机)\\n";
        }

        auto motors = dog.get_motor_status();
        int online = 0;
        for (const auto& m : motors) if (m.online) ++online;
        std::cout << "[OK] 电机在线: " << online << "/"
                  << motors.size() << "\\n";

        auto profile = dog.get_profile();
        std::cout << "[OK] 当前策略: " << profile.current << "\\n";

        std::cout << "\\n--- 连接验证通过 ---\\n";

    } catch (const orix::ConnectionError& e) {
        std::cerr << "\\n[FAIL] 连接失败: " << e.what() << "\\n";
        std::cerr << "排查步骤:\\n";
        std::cerr << "  1. 确认机器人已上电\\n";
        std::cerr << "  2. ping 192.168.66.190\\n";
        std::cerr << "  3. 检查防火墙是否阻断 13145/TCP\\n";
        return 1;
    }
    return 0;
}\
""")
    doc.add_paragraph("编译并运行：")
    code_block(doc, """\
g++ -std=c++17 -o verify_connection verify_connection.cpp \\
    -I/path/to/orix/include -L/path/to/orix/lib -lorix -lgrpc++ -lprotobuf
./verify_connection\
""")
    doc.add_paragraph()

    set_heading(doc, "2.7 第一次完整操作", 2)
    doc.add_paragraph(
        "确认连接验证通过后，尝试完整的站起--行走--坐下流程。这个流程包含了"
        "状态切换等待、停步和安全收尾，是后续编写演示程序的推荐模板。"
    )
    code_block(doc, """\
#include "orix/orix.h"
#include <iostream>
#include <thread>
#include <chrono>

int main() {
    orix::OrixClient dog("192.168.66.190");

    dog.enable();
    std::cout << "状态: " << orix::to_string(dog.get_state())
              << "\\n"; // grounded

    dog.stand_up();
    std::this_thread::sleep_for(std::chrono::seconds(4));
    std::cout << "状态: " << orix::to_string(dog.get_state())
              << "\\n"; // standing

    dog.walk(0.3);
    std::this_thread::sleep_for(std::chrono::seconds(3));

    dog.walk();       // 无参数 = 停止
    std::this_thread::sleep_for(std::chrono::seconds(1));

    dog.sit_down();
    std::this_thread::sleep_for(std::chrono::seconds(4));
    std::cout << "状态: " << orix::to_string(dog.get_state())
              << "\\n"; // grounded

    dog.disable();
    return 0;  // dog 析构自动释放 gRPC channel
}\
""")
    warning_box(doc,
        "首次操作建议在台架上进行，或有人在旁边准备随时用遥控器接管。"
        "遥控器具有最高优先级，任何时候拨动遥控器都可以立即接管控制权。"
    )
    doc.add_paragraph()

    set_heading(doc, "2.8 RAII 与资源管理", 2)
    doc.add_paragraph(
        "Python SDK 使用 with 语句管理资源，C++ SDK 使用 RAII（Resource Acquisition "
        "Is Initialization）模式。OrixClient 对象在析构时自动释放 gRPC 通道资源，无需手动调用 close()。"
    )
    code_block(doc, """\
// Python 写法：
// with OrixClient("192.168.66.190") as dog:
//     state = dog.get_state()

// C++ 写法：OrixClient 是 RAII 对象，析构自动释放
{
    orix::OrixClient dog("192.168.66.190");
    auto state = dog.get_state();
    std::cout << "状态: " << orix::to_string(state) << "\\n";
}  // dog 析构，gRPC channel 自动关闭\
""")
    tip_box(doc, "在生产代码中，建议通过 std::unique_ptr 管理 OrixClient 的生命周期，或将其作为类成员，利用 RAII 确保资源安全释放。")
    doc.add_paragraph()

    set_heading(doc, "2.9 常见问题", 2)

    set_heading(doc, "2.9.1 连接超时（DEADLINE_EXCEEDED）", 3)
    doc.add_paragraph("如果程序抛出 orix::TimeoutError，通常是网络配置或服务未启动导致的。排查步骤：")
    numbered(doc, "确认机器人已上电且启动完成")
    numbered(doc, "ping 192.168.66.190 确认网络连通")
    numbered(doc, "SSH 登录确认服务运行：ssh sunrise@192.168.66.190 'systemctl status brainstem'")
    numbered(doc, "检查防火墙是否阻断了 13145/TCP")

    set_heading(doc, "2.9.2 walk() 没反应", 3)
    doc.add_paragraph("必须按照状态机顺序操作：enable() -> stand_up() -> 等待站稳 -> walk()。")
    code_block(doc, """\
auto state = dog.get_state();
std::cout << orix::to_string(state) << "\\n";
// 如果不是 "standing"，walk() 不会生效\
""")

    set_heading(doc, "2.9.3 编译失败：找不到 gRPC", 3)
    code_block(doc, """\
cmake .. -DCMAKE_PREFIX_PATH=/path/to/grpc/install
# 或使用 vcpkg toolchain
cmake .. -DCMAKE_TOOLCHAIN_FILE=/path/to/vcpkg/scripts/buildsystems/vcpkg.cmake\
""")

    set_heading(doc, "2.9.4 操作被拒绝", 3)
    doc.add_paragraph("当状态机正在执行过渡动画时，新的运动指令会被拒绝。")
    code_block(doc, """\
try {
    dog.walk(0.3);
} catch (const orix::InvalidStateError& e) {
    std::cerr << "状态错误: " << e.what() << "\\n";
    // 状态机正在过渡中，请等待过渡完成
}\
""")
    doc.add_paragraph()

    # ================================================================
    #  第 3 章：理论基础
    # ================================================================
    set_heading(doc, "第三章 理论基础", 1)

    set_heading(doc, "3.1 PD 位置控制", 2)
    doc.add_paragraph(
        "四足机器人的每个关节电机均采用 PD（比例-微分）位置控制器。PD 控制器根据"
        "目标位置与当前位置的偏差，以及目标速度与当前速度的偏差，计算驱动关节所需的力矩。"
    )
    doc.add_paragraph("基本控制律：")
    code_block(doc, "tau = Kp * (q_target - q_current) + Kd * (dq_target - dq_current)")
    doc.add_paragraph("其中：")
    bullet(doc, "tau --- 关节输出力矩（N*m）")
    bullet(doc, "Kp --- 比例增益（位置刚度），单位 N*m/rad")
    bullet(doc, "Kd --- 微分增益（速度阻尼），单位 N*m*s/rad")
    bullet(doc, "q_target, q_current --- 目标与当前关节角度（rad）")
    bullet(doc, "dq_target, dq_current --- 目标与当前关节角速度（rad/s）")

    set_heading(doc, "3.1.1 物理意义", 3)
    doc.add_paragraph(
        "Kp 决定了关节的刚度：Kp 越大，关节对位置偏差的响应越强，如同弹簧越硬。"
        "Kd 决定了关节的阻尼：Kd 越大，关节运动越平滑，振荡越少，如同液压阻尼器。"
    )

    set_heading(doc, "3.1.2 稳定性条件", 3)
    doc.add_paragraph("PD 控制器稳定的必要条件为：Kp > 0, Kd > 0。")

    set_heading(doc, "3.1.3 阻尼比分析", 3)
    doc.add_paragraph(
        "将单关节建模为二阶系统 I*ddq + Kd*dq + Kp*q = 0（I 为关节转动惯量），"
        "可定义阻尼比和固有频率。不同阻尼比对应不同的动态行为："
    )
    bullet(doc, "zeta < 1（欠阻尼）：关节到达目标位置时会出现振荡，适用于需要快速响应的场景")
    bullet(doc, "zeta = 1（临界阻尼）：关节以最快速度到达目标位置且不超调，理论最优")
    bullet(doc, "zeta > 1（过阻尼）：关节缓慢趋近目标位置，无超调但响应速度慢")
    doc.add_paragraph("在实际机器人中，通常选取略欠阻尼（zeta ~ 0.7-0.9）的参数，以兼顾响应速度和稳定性。")

    set_heading(doc, "3.2 强化学习运动控制", 2)
    doc.add_paragraph(
        "ORIX 采用基于强化学习（RL）的方法训练四足步态策略。策略网络在高保真物理仿真器"
        "（Isaac Lab）中通过数百万次交互学习，然后直接部署到真实机器人上（Sim-to-Real）。"
    )

    set_heading(doc, "3.2.1 策略网络", 3)
    doc.add_paragraph("策略网络 pi 是一个将观测映射到动作的神经网络：a_t = pi(o_t)")

    set_heading(doc, "3.2.2 观测空间", 3)
    doc.add_paragraph("单帧观测向量维度为 57，由以下部分组成：")
    add_table(doc, ["分量", "维度", "说明"], [
        ["omega（角速度）", "3", "陀螺仪角速度（机体坐标系）"],
        ["g_proj（投影重力）", "3", "投影重力向量"],
        ["c_cmd（速度指令）", "3", "速度指令 (vx, vy, omega_z)"],
        ["q_pos（关节位置）", "12", "关节位置（12 腿关节，rad）"],
        ["dq（关节速度）", "12", "关节速度（rad/s）"],
        ["a_prev（上一动作）", "12", "上一时刻动作输出"],
        ["b_contact（接触）", "12", "足端接触布尔量"],
    ])

    set_heading(doc, "3.2.3 动作空间与缩放", 3)
    doc.add_paragraph(
        "策略关节位置观测以 Profile 的 policyDefaultPose 为训练零点，策略动作也必须用"
        "同一个零点反归一化：q_relative_i = q_actual_i - q_policyDefault_i；"
        "q_target_i = q_policyDefault_i + a_i * actionScale_i。"
    )
    doc.add_paragraph(
        "policyDefaultPose 是策略训练接口的一部分，不一定等于安全站起时的物理姿态 "
        "standUpPose。例如，某腿的 standUpPose 可为 [0.1, -0.8, 1.8]，而该腿的 "
        "policyDefaultPose 仍使用训练值 [0.0, -1.1, 2.6]。"
    )

    set_heading(doc, "3.2.4 Sim-to-Real", 3)
    doc.add_paragraph(
        "策略在 Isaac Lab 仿真环境中训练，使用域随机化（Domain Randomization）缩小"
        "仿真与现实的差距。训练完成后，策略网络导出为 ONNX 格式，部署到真实机器人的"
        "嵌入式处理器上进行实时推理。"
    )

    set_heading(doc, "3.3 姿态表示与坐标变换", 2)

    set_heading(doc, "3.3.1 四元数基础", 3)
    doc.add_paragraph(
        "单位四元数 q = (w, x, y, z)，||q|| = sqrt(w^2 + x^2 + y^2 + z^2) = 1。"
        "机器人的姿态（朝向）在三维空间中用四元数表示，避免了欧拉角的万向锁问题。"
    )

    set_heading(doc, "3.3.2 投影重力计算", 3)
    doc.add_paragraph(
        "投影重力（Projected Gravity）是将世界坐标系的重力方向变换到机体坐标系，"
        "用于让策略感知机器人当前的倾斜状态：g_body = R^T * [0, 0, -1]^T"
    )

    set_heading(doc, "3.3.3 Hamilton 约定与 ROS 约定", 3)
    bullet(doc, "Hamilton 约定（ORIX / Protobuf）：(w, x, y, z) --- 实部 w 在前")
    bullet(doc, "ROS 约定（ROS2 / tf2）：(x, y, z, w) --- 实部 w 在后")
    doc.add_paragraph("在 ORIX 系统内部统一使用 Hamilton 约定。当与 ROS 系统交互时，需手动调整字段顺序。")

    set_heading(doc, "3.4 有限状态机", 2)
    doc.add_paragraph(
        "ORIX 的运动控制流程由一个有限状态机（FSM）管理。FSM 确保机器人在任何时刻都"
        "处于明确定义的状态，且所有状态转换都经过安全校验。"
    )
    doc.add_paragraph("FSM 形式化定义为五元组：FSM = (S, Sigma, delta, s0, F)")
    doc.add_paragraph("S = {Zero, Grounded, Standing, Walking, Transitioning}")
    doc.add_paragraph("Sigma = {init, standUp, sitDown, walk, stop, fault}")
    doc.add_paragraph("s0 = Zero")
    doc.add_paragraph(
        "安全不变量：从任意状态出发，收到 fault 信号后，"
        "系统必须经由安全坐下动画回到 Grounded 状态。"
    )

    set_heading(doc, "3.5 控制频率与实时性", 2)
    doc.add_paragraph(
        "ORIX 的主控制循环以 50 Hz 固定频率运行，即每个控制周期为 T = 1/50 = 20ms。"
    )
    doc.add_paragraph(
        "C++ SDK 相较于 Python SDK 的一大优势在于：C++ 没有 GIL 限制，可以在多线程中"
        "真正并行执行计算密集型任务。这使得 C++ SDK 特别适合以下场景："
    )
    bullet(doc, "在独立线程中以 50Hz 订阅 IMU 数据，同时在主线程中执行运动控制")
    bullet(doc, "将传感器数据处理和运动指令发送完全解耦")
    bullet(doc, "在机器人本体上直接运行，消除网络延迟")
    doc.add_paragraph()

    # ================================================================
    #  第 4 章：硬件概览
    # ================================================================
    set_heading(doc, "第四章 硬件概览", 1)

    set_heading(doc, "4.1 整体架构", 2)
    add_table(doc, ["组件", "型号/规格", "说明"], [
        ["主控板", "地瓜机器人 S100P（RDK X5）", "AI 算力 128 TOPS (Nash BPU)"],
        ["IMU", "HI91 惯性测量单元", "串口 /dev/ttyUSB1，约 50Hz"],
        ["电机", "Robstride 伺服电机 x16", "CAN 总线通信"],
        ["电池", "7 串锂电池", "正常 24V，低压保护 18V"],
        ["通信", "RJ45 以太网 + WiFi", "gRPC 端口 13145"],
        ["遥控器", "YUNZHUO 无线遥控器", "最高控制优先级"],
    ])

    set_heading(doc, "4.2 关键参数", 2)
    add_table(doc, ["参数", "值"], [
        ["自由度", "12 DOF（每腿 3 旋转关节）"],
        ["控制频率", "50 Hz (20ms 周期)"],
        ["推理引擎", "ONNX Runtime"],
        ["通信协议", "gRPC over HTTP/2"],
        ["默认 IP", "192.168.66.190"],
        ["gRPC 端口", "13145"],
        ["SSH 端口", "22（用户 sunrise / 密码 sunrise）"],
        ["正常电压", "约 24V"],
        ["低压保护", "18V"],
    ])

    set_heading(doc, "4.3 CAN 总线拓扑", 2)
    doc.add_paragraph(
        "16 个 Robstride 伺服电机通过 4 路 CAN 总线与主控板通信，每路 CAN 总线连接"
        "一条腿的 4 个电机。CAN 通信由 brainstem 控制栈在底层管理，SDK 用户无需直接操作 CAN 总线。"
    )

    set_heading(doc, "4.4 传感器概览", 2)
    add_table(doc, ["传感器", "位置", "接口", "数据"], [
        ["IMU (HI91)", "机身中心", "串口", "角速度 + 四元数"],
        ["关节编码器", "每个电机内置", "CAN", "位置/速度/力矩"],
        ["电压传感器", "每个电机内置", "CAN", "母线电压"],
        ["温度传感器", "每个电机内置", "CAN", "电机温度"],
    ])
    doc.add_paragraph()

    # ================================================================
    #  第 5 章：坐标系与关节
    # ================================================================
    set_heading(doc, "第五章 坐标系与关节", 1)

    set_heading(doc, "5.1 机体坐标系", 2)
    doc.add_paragraph("ORIX 使用右手坐标系，原点位于机身几何中心：")
    add_table(doc, ["轴", "正方向", "说明"], [
        ["X", "前方 (Forward)", "机器人头部方向"],
        ["Y", "左方 (Left)", "机器人左侧方向"],
        ["Z", "上方 (Up)", "垂直向上"],
    ])
    doc.add_paragraph()
    doc.add_paragraph("walk() 指令与坐标系的关系：")
    add_table(doc, ["参数", "正值", "负值", "对应轴"], [
        ["vx", "前进", "后退", "+X / -X"],
        ["vy", "向左", "向右", "+Y / -Y"],
        ["vyaw", "逆时针", "顺时针", "绕 +Z 轴"],
    ])

    set_heading(doc, "5.2 关节编号总表", 2)
    doc.add_paragraph("ORIX 全部 12 个关节的编号和 URDF 名称：")
    add_table(doc, ["索引", "名称", "腿", "关节类型", "URDF 名称"], [
        ["0",  "FR Hip",   "前右", "髋 (横滚)",   "fr_hip_joint"],
        ["1",  "FR Thigh", "前右", "大腿 (俯仰)", "fr_thigh_joint"],
        ["2",  "FR Calf",  "前右", "小腿 (俯仰)", "fr_calf_joint"],
        ["3",  "FL Hip",   "前左", "髋 (横滚)",   "fl_hip_joint"],
        ["4",  "FL Thigh", "前左", "大腿 (俯仰)", "fl_thigh_joint"],
        ["5",  "FL Calf",  "前左", "小腿 (俯仰)", "fl_calf_joint"],
        ["6",  "RR Hip",   "后右", "髋 (横滚)",   "rr_hip_joint"],
        ["7",  "RR Thigh", "后右", "大腿 (俯仰)", "rr_thigh_joint"],
        ["8",  "RR Calf",  "后右", "小腿 (俯仰)", "rr_calf_joint"],
        ["9",  "RL Hip",   "后左", "髋 (横滚)",   "rl_hip_joint"],
        ["10", "RL Thigh", "后左", "大腿 (俯仰)", "rl_thigh_joint"],
        ["11", "RL Calf",  "后左", "小腿 (俯仰)", "rl_calf_joint"],
    ])

    set_heading(doc, "5.3 关节分组", 2)

    set_heading(doc, "5.3.1 按腿分组", 3)
    add_table(doc, ["腿", "关节索引"], [
        ["前右 (FR)", "[0, 1, 2]"],
        ["前左 (FL)", "[3, 4, 5]"],
        ["后右 (RR)", "[6, 7, 8]"],
        ["后左 (RL)", "[9, 10, 11]"],
    ])

    set_heading(doc, "5.3.2 按关节类型分组", 3)
    add_table(doc, ["关节类型", "索引", "功能"], [
        ["所有 Hip",   "[0, 3, 6, 9]",   "髋关节，控制腿的内外摆动"],
        ["所有 Thigh", "[1, 4, 7, 10]",  "大腿关节，控制腿的前后摆动"],
        ["所有 Calf",  "[2, 5, 8, 11]",  "小腿关节，控制膝盖弯曲"],
    ])

    set_heading(doc, "5.4 典型姿态", 2)
    doc.add_paragraph(
        "Profile 使用 standUpPose 表示站起、回站和 gesture 的物理目标，使用 "
        "policyDefaultPose 表示策略关节相对观测和动作反归一化的训练零点。"
    )
    doc.add_paragraph(
        "旧 standingPose 仅用于向后兼容：缺少 standUpPose 时，standUpPose 单独回退到 "
        "standingPose；缺少 policyDefaultPose 时，policyDefaultPose 单独回退到 "
        "standingPose。因此只含 standingPose 的旧 Profile 行为不变。"
    )
    code_block(doc, """\
// 同一条腿的示例；实际 Profile 为完整关节数组
standUpPose:       [0.1, -0.8, 1.8]  // 物理安全站姿
policyDefaultPose: [0.0, -1.1, 2.6]  // ONNX 训练零点\
""")
    warning_box(
        doc,
        "不得用 standUpPose 替换策略零点，否则 observation 和动作目标都会偏离 ONNX 训练接口。",
    )

    set_heading(doc, "5.5 PD 控制参数", 2)
    doc.add_paragraph("推理阶段（行走时）的 PD 参数：")
    add_table(doc, ["索引", "关节", "Kp", "Kd", "actionScale"], [
        ["0, 3, 6, 9",   "Hip",   "80",  "10", "0.15"],
        ["1, 4, 7, 10",  "Thigh", "100", "10", "0.25"],
        ["2, 5, 8, 11",  "Calf",  "120", "15", "0.25"],
    ])

    set_heading(doc, "5.6 编程中的关节操作", 2)

    set_heading(doc, "5.6.1 读取关节数据", 3)
    code_block(doc, """\
// 实时流式读取关节数据
dog.listen_joint([](const orix::JointData& j) -> bool {
    std::cout << "关节 " << j.id
              << ": pos=" << j.position << "rad"
              << " vel=" << j.velocity << "rad/s"
              << " trq=" << j.torque << "Nm\\n";
    return true;  // 继续接收
});\
""")

    set_heading(doc, "5.6.2 诊断单条腿", 3)
    code_block(doc, """\
// 检查前右腿 (FR) 的 4 个电机
auto motors = dog.get_motor_status();
for (const auto& m : motors) {
    if (m.id >= 0 && m.id <= 3) {  // FR 腿
        std::cout << "  关节" << m.id
                  << ": online=" << m.online
                  << " temp=" << m.temperature << "C\\n";
    }
}\
""")

    set_heading(doc, "5.6.3 清除单腿故障", 3)
    code_block(doc, """\
// 只清除前右腿的电机故障
dog.clear_motor_fault({0, 1, 2, 3});\
""")
    doc.add_paragraph()

    # ================================================================
    #  第 6 章：状态机与运动控制
    # ================================================================
    set_heading(doc, "第六章 状态机与运动控制", 1)

    set_heading(doc, "6.1 状态机总览", 2)
    doc.add_paragraph(
        "ORIX 的运动控制由一个严格的有限状态机管理。所有 SDK 指令必须遵循状态机规则，"
        "非法的状态转换会被拒绝并抛出 orix::InvalidStateError 异常。"
    )

    set_heading(doc, "6.1.1 五个状态", 3)
    add_table(doc, ["状态", "SDK 返回值", "说明"], [
        ["Zero",          "RobotState::Zero",          "上电初始态，FSM 尚未初始化"],
        ["Grounded",      "RobotState::Grounded",      "坐姿/趴地，安全状态"],
        ["Standing",      "RobotState::Standing",      "直立站稳，可接收行走指令"],
        ["Walking",       "RobotState::Walking",       "行走中，RL 策略以 50Hz 推理"],
        ["Transitioning", "RobotState::Transitioning", "过渡中，拒绝新的运动指令"],
    ])

    set_heading(doc, "6.1.2 状态转换路径", 3)
    add_table(doc, ["起始状态", "触发条件", "中间状态", "目标状态"], [
        ["Zero",     "系统初始化",         "--",           "Grounded"],
        ["Grounded", "stand_up()",         "Transitioning","Standing"],
        ["Standing", "walk(vx, vy, vyaw)", "--",           "Walking"],
        ["Walking",  "walk() 停步",        "Transitioning","Standing"],
        ["Standing", "sit_down()",         "Transitioning","Grounded"],
        ["Walking",  "sit_down()",         "先停再坐",     "Grounded"],
        ["任意状态", "Fault 故障",          "自动过渡",     "Grounded"],
    ])

    set_heading(doc, "6.2 状态转换详解", 2)

    set_heading(doc, "6.2.1 Grounded -> Standing（站起）", 3)
    add_table(doc, ["项目", "说明"], [
        ["触发", "dog.stand_up()"],
        ["耗时", "150 ticks x 20ms = 3 秒"],
        ["过程", "从当前姿态线性插值到 standUpPose"],
    ])

    set_heading(doc, "6.2.2 Standing -> Walking（开始行走）", 3)
    add_table(doc, ["项目", "说明"], [
        ["触发", "dog.walk(vx, vy, vyaw)"],
        ["耗时", "即时切换"],
        ["过程", "RL 策略接管，以 50Hz 推理输出关节指令"],
    ])

    set_heading(doc, "6.2.3 Standing -> Grounded（坐下）", 3)
    add_table(doc, ["项目", "说明"], [
        ["触发", "dog.sit_down()"],
        ["耗时", "150 ticks x 20ms = 3 秒"],
        ["过程", "从 standUpPose 线性插值到 sittingPose (全零)"],
    ])

    set_heading(doc, "6.3 非法操作处理", 2)
    add_table(doc, ["当前状态", "发送的指令", "结果"], [
        ["Grounded",      "walk()",    "抛出 InvalidStateError"],
        ["Grounded",      "sit_down()","忽略（已在坐姿）"],
        ["Standing",      "stand_up()","忽略（已站立）"],
        ["Transitioning", "任何运动指令","抛出 InvalidStateError"],
        ["Zero",          "任何运动指令","忽略"],
    ])
    doc.add_paragraph("最佳实践：在发送指令前先查询状态。")
    code_block(doc, """\
auto state = dog.get_state();
if (state == orix::RobotState::Standing) {
    dog.walk(0.3);
} else if (state == orix::RobotState::Grounded) {
    dog.stand_up();
}\
""")

    set_heading(doc, "6.4 walk() 速度约定", 2)
    code_block(doc, "dog.walk(vx, vy, vyaw);")
    add_table(doc, ["参数", "范围", "正值", "负值", "零值"], [
        ["vx",   "[-1.0, 1.0]", "前进 (+X)", "后退 (-X)", "不动"],
        ["vy",   "[-1.0, 1.0]", "向左 (+Y)", "向右 (-Y)", "不动"],
        ["vyaw", "[-1.0, 1.0]", "逆时针 (+Z)","顺时针 (-Z)", "不转"],
    ])
    code_block(doc, """\
dog.walk(0.5);                 // 前进
dog.walk(-0.3);                // 后退
dog.walk(0.0, 0.3);           // 向左平移
dog.walk(0.0, 0.0, 0.3);     // 原地逆时针转
dog.walk(0.3, 0.0, 0.2);     // 前进 + 左转
dog.walk();                    // 停步\
""")

    set_heading(doc, "6.5 电机使能/禁用", 2)
    doc.add_paragraph("enable() 使能全部 16 个电机，电机通电进入闭环位置控制。")
    doc.add_paragraph("disable() 禁用全部电机（安全停止），电机断电，关节自由落体。建议在 Grounded 状态下调用。")

    set_heading(doc, "6.6 控制权仲裁 (ControlArbiter)", 2)
    doc.add_paragraph("ORIX 支持多个控制源同时存在，通过 ControlArbiter 仲裁控制权。遥控器始终具有最高优先级。")
    add_table(doc, ["场景", "行为"], [
        ["SDK 控制中，遥控器摇杆动",   "SDK 被立即抢占，遥控器接管"],
        ["遥控器控制中，SDK 发指令",    "SDK 指令被忽略"],
        ["遥控器松手不动",              "3 秒超时后自动释放，SDK 恢复控制权"],
        ["SDK 控制中，遥控器关机",      "SDK 保持控制权，不受影响"],
    ])

    set_heading(doc, "6.7 故障处理", 2)
    doc.add_paragraph("检测到 Fault 信号后，FSM 自动执行安全坐下动画回到 Grounded 状态。")

    set_heading(doc, "6.8 策略切换", 2)
    doc.add_paragraph("策略切换必须在 Grounded 状态下执行：")
    code_block(doc, """\
dog.sit_down();
std::this_thread::sleep_for(std::chrono::seconds(4));
auto new_profile = dog.switch_profile("mini");
std::cout << "已切换到: " << new_profile.current << "\\n";\
""")

    set_heading(doc, "6.9 典型使用流程", 2)

    set_heading(doc, "6.9.1 标准操作序列", 3)
    code_block(doc, """\
#include "orix/orix.h"
#include <algorithm>
#include <chrono>
#include <iostream>
#include <thread>

int main() {
    orix::OrixClient dog("192.168.66.190");

    // 1. 诊断检查
    auto state = dog.get_state();
    std::cout << "初始状态: " << orix::to_string(state) << "\\n";

    auto voltages = dog.get_voltage();
    double avg_v = 0;
    for (auto v : voltages) avg_v += v;
    avg_v /= voltages.size();
    std::cout << "电池电压: " << avg_v << "V\\n";

    auto motors = dog.get_motor_status();
    int offline = 0;
    for (const auto& m : motors) if (!m.online) ++offline;
    std::cout << "电机在线: " << (16 - offline) << "/16\\n";

    // 2. 使能电机
    dog.enable();

    // 3. 站起 (3 秒过渡)
    dog.stand_up();
    std::this_thread::sleep_for(std::chrono::seconds(4));

    // 4. 行走
    dog.walk(0.3);
    std::this_thread::sleep_for(std::chrono::seconds(3));
    dog.walk(0.3, 0.0, 0.2);   // 前进 + 左转
    std::this_thread::sleep_for(std::chrono::seconds(2));
    dog.walk();                 // 停步

    // 5. 坐下 (3 秒过渡)
    dog.sit_down();
    std::this_thread::sleep_for(std::chrono::seconds(4));

    // 6. 禁用电机
    dog.disable();
    return 0;
}\
""")

    set_heading(doc, "6.9.2 实时数据监控", 3)
    code_block(doc, """\
#include "orix/orix.h"
#include <atomic>
#include <chrono>
#include <iostream>
#include <thread>

int main() {
    orix::OrixClient dog("192.168.66.190");
    std::atomic<bool> running{true};

    // 后台线程监听 IMU
    std::thread imu_thread([&] {
        dog.listen_imu([&](const orix::ImuData& d) -> bool {
            std::cout << "\\rIMU: gx=" << d.gyroscope.x
                      << " gy=" << d.gyroscope.y
                      << " gz=" << d.gyroscope.z << std::flush;
            return running.load();
        });
    });

    // 后台线程监听状态变化
    std::thread state_thread([&] {
        dog.listen_state([&](orix::RobotState s) -> bool {
            std::cout << "\\n[STATE] -> " << orix::to_string(s) << "\\n";
            return running.load();
        });
    });

    // 主线程执行运动
    dog.enable();
    dog.stand_up();
    std::this_thread::sleep_for(std::chrono::seconds(4));
    dog.walk(0.3);
    std::this_thread::sleep_for(std::chrono::seconds(5));
    dog.sit_down();
    std::this_thread::sleep_for(std::chrono::seconds(4));
    dog.disable();

    running = false;
    imu_thread.join();
    state_thread.join();
    return 0;
}\
""")

    set_heading(doc, "6.9.3 安全包装器", 3)
    code_block(doc, """\
void safe_run(orix::OrixClient& dog) {
    dog.enable();
    dog.stand_up();
    std::this_thread::sleep_for(std::chrono::seconds(4));
    dog.walk(0.3);
    std::this_thread::sleep_for(std::chrono::seconds(5));
}

int main() {
    orix::OrixClient dog("192.168.66.190");
    try {
        safe_run(dog);
    } catch (const std::exception& e) {
        // 任何异常都执行安全关机
    }
    try { dog.walk();    } catch (...) {}
    try { dog.sit_down(); } catch (...) {}
    for (int i = 0; i < 50; ++i) {
        if (dog.get_state() == orix::RobotState::Grounded) break;
        std::this_thread::sleep_for(std::chrono::milliseconds(200));
    }
    try { dog.disable(); } catch (...) {}
    return 0;
}\
""")

    set_heading(doc, "6.10 关键参数速查", 2)
    add_table(doc, ["参数", "值", "说明"], [
        ["控制频率",   "50 Hz (20ms)",    "主控制循环和推理频率"],
        ["gRPC 端口",  "13145",           "SDK 连接端口"],
        ["站起耗时",   "3 秒 (150 ticks)","Grounded -> Standing"],
        ["坐下耗时",   "3 秒 (150 ticks)","Standing -> Grounded"],
        ["仲裁超时",   "3 秒",            "遥控器释放控制权等待时间"],
        ["低压阈值",   "18V",             "低于此值触发保护坐下"],
        ["启动超时",   "10 秒",           "FSM 等待 Grounded 超时"],
        ["关机超时",   "8 秒",            "安全关机等待超时"],
        ["RPC 超时",   "5 秒",            "SDK 默认 RPC 超时"],
    ])
    doc.add_paragraph()

    # ================================================================
    #  第 7 章：实用开发指南
    # ================================================================
    set_heading(doc, "第七章 实用开发指南", 1)

    set_heading(doc, "7.1 电池管理", 2)
    doc.add_paragraph("机器狗采用 7 串锂电池，正常工作电压约 24V，低压保护阈值 18V。在长时间运行的应用中，建议实现主动电池监控。")
    code_block(doc, """\
#include "orix/orix.h"
#include <iostream>
#include <algorithm>

double estimate_battery_percent(orix::OrixClient& dog) {
    auto voltages = dog.get_voltage();
    if (voltages.empty()) return -1.0;

    double sum = 0;
    int count = 0;
    for (auto v : voltages) {
        if (v > 0) { sum += v; ++count; }
    }
    if (count == 0) return -1.0;

    double avg_v = sum / count;
    // 线性插值: 29.4V=100%, 18V=0%
    double pct = std::clamp((avg_v - 18.0) / (29.4 - 18.0) * 100.0,
                            0.0, 100.0);
    return pct;
}

int main() {
    orix::OrixClient dog("192.168.66.190");
    double pct = estimate_battery_percent(dog);
    auto voltages = dog.get_voltage();
    double min_v = *std::min_element(voltages.begin(), voltages.end());

    std::cout << "电量: " << pct << "% (" << min_v << "V)\\n";
    if (pct < 30.0) {
        std::cout << "低电量，建议收工充电。\\n";
    }
    return 0;
}\
""")

    set_heading(doc, "7.2 数据记录与采集", 2)
    doc.add_paragraph("记录 IMU 数据到 CSV：")
    code_block(doc, """\
#include "orix/orix.h"
#include <chrono>
#include <fstream>
#include <iostream>

int main() {
    orix::OrixClient dog("192.168.66.190");
    std::ofstream csv("imu_log.csv");
    csv << "timestamp_ms,gx,gy,gz,qw,qx,qy,qz\\n";

    int count = 0;
    auto t0 = std::chrono::steady_clock::now();

    dog.listen_imu([&](const orix::ImuData& d) -> bool {
        csv << d.timestamp_ms   << ","
            << d.gyroscope.x  << "," << d.gyroscope.y
            << "," << d.gyroscope.z  << ","
            << d.quaternion.w << "," << d.quaternion.x
            << "," << d.quaternion.y << ","
            << d.quaternion.z << "\\n";
        ++count;
        auto elapsed = std::chrono::steady_clock::now() - t0;
        return elapsed < std::chrono::seconds(10);
    });

    std::cout << "记录完成：" << count << " 帧 -> imu_log.csv\\n";
    return 0;
}\
""")

    set_heading(doc, "7.3 多线程流式订阅", 2)
    doc.add_paragraph(
        "C++ SDK 的流式接口（listen_imu、listen_joint、listen_state）是阻塞调用，"
        "建议在独立线程中运行。C++ 没有 GIL 限制，多线程可以真正并行执行。"
    )
    code_block(doc, """\
#include "orix/orix.h"
#include <atomic>
#include <chrono>
#include <iostream>
#include <thread>

int main() {
    orix::OrixClient dog("192.168.66.190");
    std::atomic<bool> running{true};

    std::thread imu_thread([&] {
        dog.listen_imu([&](const orix::ImuData& d) -> bool {
            std::cout << "gyro_x=" << d.gyroscope.x
                      << " qw=" << d.quaternion.w
                      << " t=" << d.timestamp_ms << "ms\\n";
            return running.load();
        });
    });

    std::thread state_thread([&] {
        dog.listen_state([&](orix::RobotState s) -> bool {
            std::cout << "[STATE] -> " << orix::to_string(s) << "\\n";
            return running.load();
        });
    });

    // 主线程执行运动控制
    dog.enable();
    dog.stand_up();
    std::this_thread::sleep_for(std::chrono::seconds(4));
    dog.walk(0.3);
    std::this_thread::sleep_for(std::chrono::seconds(5));
    dog.sit_down();
    std::this_thread::sleep_for(std::chrono::seconds(4));
    dog.disable();

    running = false;
    imu_thread.join();
    state_thread.join();
    return 0;
}\
""")
    tip_box(doc, "使用 std::atomic<bool> 控制流的生命周期，是 C++ 多线程流式订阅的标准模式。callback 返回 false 时流会自动停止。")

    set_heading(doc, "7.4 集成 OpenCV 摄像头", 2)
    code_block(doc, """\
#include "orix/orix_camera.h"
#include <iostream>

int main() {
    orix::OrixCamera cam(0, 640, 480, 30);
    cam.open();

    cam.save_photo("photo.jpg");
    std::cout << "照片已保存: photo.jpg\\n";

    cam.start_recording("output.mp4", 10.0);
    std::cout << "正在录制 10 秒...\\n";
    std::this_thread::sleep_for(std::chrono::seconds(11));

    cam.stream_mjpeg(8080);
    std::cout << "MJPEG 流: http://localhost:8080/stream\\n";
    std::this_thread::sleep_for(std::chrono::seconds(30));

    cam.stop_stream();
    cam.close();
    return 0;
}\
""")

    set_heading(doc, "7.4.1 视觉引导行走", 3)
    doc.add_paragraph("结合 OpenCV 实现基于颜色追踪的行走控制：")
    code_block(doc, """\
#include "orix/orix.h"
#include "orix/orix_camera.h"
#include <opencv2/opencv.hpp>

int main() {
    orix::OrixClient dog("192.168.66.190");
    orix::OrixCamera cam(0, 640, 480, 30);
    cam.open();
    dog.enable();
    dog.stand_up();
    std::this_thread::sleep_for(std::chrono::seconds(4));

    try {
        for (int i = 0; i < 500; ++i) {
            cv::Mat frame = cam.read();
            if (frame.empty()) continue;
            cv::Mat hsv;
            cv::cvtColor(frame, hsv, cv::COLOR_BGR2HSV);
            cv::Mat mask;
            cv::inRange(hsv, cv::Scalar(0, 100, 100),
                        cv::Scalar(10, 255, 255), mask);
            auto moments = cv::moments(mask);
            if (moments.m00 > 500) {
                int cx = static_cast<int>(moments.m10 / moments.m00);
                int center = frame.cols / 2;
                double turn = -(cx - center)
                              / static_cast<double>(center) * 0.3;
                dog.walk(0.2, 0.0, turn);
            } else {
                dog.walk();
            }
            std::this_thread::sleep_for(std::chrono::milliseconds(30));
        }
    } catch (...) {}

    dog.walk();
    dog.sit_down();
    std::this_thread::sleep_for(std::chrono::seconds(4));
    dog.disable();
    cam.close();
    return 0;
}\
""")

    set_heading(doc, "7.5 自定义 Profile", 2)
    doc.add_paragraph("Profile 是机器狗的运动参数集合。自定义 Profile 的 JSON 文件结构：")
    code_block(doc, """\
{
  "name": "my_custom",
  "description": "自定义策略",
  "modelPath": "model/my_policy.onnx",
  "standingPose": [0.0, -0.65, 1.5, 0.0, ...],
  "standUpPose": [0.1, -0.8, 1.8, 0.0, ...],
  "policyDefaultPose": [0.0, -1.1, 2.6, 0.0, ...],
  "sittingPose": [0.0, 0.0, 0.0, 0.0, ...],
  "inferKp": [...], "inferKd": [...],
  "standUpKp": [...], "standUpKd": [...],
  "actionScale": [0.125, 0.25, 0.25, 5.0]
}\
""")
    doc.add_paragraph(
        "standUpPose 和 policyDefaultPose 都是可选的兼容扩展：前者缺失时回退到 "
        "standingPose，后者缺失时也独立回退到 standingPose。建议新 Profile 显式填写"
        "两者；仅在两种语义确实相同时才给出相同数组。"
    )
    doc.add_paragraph("部署和切换：")
    code_block(doc, """\
# 上传到机器狗
scp my_custom.json sunrise@192.168.66.190:~/profiles/\
""")
    code_block(doc, """\
// 查看可用策略
auto profile = dog.get_profile();
std::cout << "当前: " << profile.current << "\\n";
for (const auto& name : profile.available) {
    std::cout << "  可用: " << name << "\\n";
}

// 切换策略（必须在 Grounded 状态）
auto new_profile = dog.switch_profile("my_custom");
std::cout << "已切换到: " << new_profile.current << "\\n";\
""")

    set_heading(doc, "7.6 CMake 工程集成最佳实践", 2)
    code_block(doc, """\
cmake_minimum_required(VERSION 3.16)
project(my_robot_app LANGUAGES CXX)

set(CMAKE_CXX_STANDARD 17)
set(CMAKE_CXX_STANDARD_REQUIRED ON)

find_package(orix REQUIRED)
find_package(gRPC REQUIRED)
find_package(Protobuf REQUIRED)
find_package(OpenCV 4.5 QUIET)  # 可选

add_executable(my_app main.cpp)
target_link_libraries(my_app PRIVATE
    orix::orix
    gRPC::grpc++
    protobuf::libprotobuf
)

if(OpenCV_FOUND)
    target_link_libraries(my_app PRIVATE
        orix::orix_camera
        ${OpenCV_LIBS}
    )
endif()\
""")
    tip_box(doc, "使用 vcpkg 作为包管理器时，在 cmake 命令中添加 -DCMAKE_TOOLCHAIN_FILE=.../vcpkg.cmake 即可自动解析所有依赖路径。")
    doc.add_paragraph()

    # ================================================================
    #  第 8 章：传感器数据
    # ================================================================
    set_heading(doc, "第八章 传感器数据", 1)

    set_heading(doc, "8.1 传感器总览", 2)
    add_table(doc, ["传感器", "硬件", "数据内容", "接口方法", "频率", "传输方式"], [
        ["IMU",  "HI91",      "角速度 + 四元数",   "listen_imu()",        "~50 Hz",  "串口 -> gRPC 流"],
        ["关节", "Robstride",  "位置/速度/力矩",    "listen_joint()",      "~50 Hz",  "CAN -> gRPC 流"],
        ["电机", "Robstride",  "在线/温度/故障",     "get_motor_status()",  "按需",    "CAN -> gRPC"],
        ["电压", "Robstride",  "16 路母线电压",      "get_voltage()",       "按需",    "CAN -> gRPC"],
        ["状态", "软件",       "FSM 状态",           "listen_state()",      "事件",    "gRPC 流"],
    ])

    set_heading(doc, "8.2 IMU 数据", 2)

    set_heading(doc, "8.2.1 数据字段", 3)
    add_table(doc, ["字段", "类型", "单位", "说明"], [
        ["gyroscope.x",  "double", "rad/s",  "绕机体 X 轴角速度 (Roll)"],
        ["gyroscope.y",  "double", "rad/s",  "绕机体 Y 轴角速度 (Pitch)"],
        ["gyroscope.z",  "double", "rad/s",  "绕机体 Z 轴角速度 (Yaw)"],
        ["quaternion.w", "double", "无量纲", "四元数实部"],
        ["quaternion.x", "double", "无量纲", "四元数虚部 i"],
        ["quaternion.y", "double", "无量纲", "四元数虚部 j"],
        ["quaternion.z", "double", "无量纲", "四元数虚部 k"],
        ["timestamp_ms", "double", "ms",     "数据时间戳"],
    ])

    set_heading(doc, "8.2.2 订阅 IMU 数据", 3)
    code_block(doc, """\
#include "orix/orix.h"
#include <iostream>
#include <iomanip>

int main() {
    orix::OrixClient dog("192.168.66.190");
    std::cout << std::fixed << std::setprecision(4);
    int count = 0;

    dog.listen_imu([&](const orix::ImuData& d) -> bool {
        ++count;
        if (count % 25 == 0) {
            std::cout << "gyro: ["
                << d.gyroscope.x << ", "
                << d.gyroscope.y << ", "
                << d.gyroscope.z << "]  "
                << "quat: w=" << d.quaternion.w
                << "  t=" << d.timestamp_ms << "ms\\n";
        }
        return count < 500;  // 约 10 秒
    });

    std::cout << "共接收 " << count << " 帧\\n";
    return 0;
}\
""")
    warning_box(doc, "ORIX SDK 使用 Hamilton 约定 (w, x, y, z)，ROS 使用 (x, y, z, w)。如需与 ROS 对接，请手动调整字段顺序。")

    set_heading(doc, "8.3 关节编码器", 2)

    set_heading(doc, "8.3.1 数据字段", 3)
    add_table(doc, ["字段", "类型", "单位", "说明"], [
        ["id",       "int",    "--",    "关节编号 (0-15)"],
        ["position", "double", "rad",   "当前角位置"],
        ["velocity", "double", "rad/s", "当前角速度"],
        ["torque",   "double", "Nm",    "当前力矩"],
        ["status",   "int",    "--",    "状态码 (0 = 正常)"],
    ])

    set_heading(doc, "8.3.2 订阅关节数据", 3)
    code_block(doc, """\
int count = 0;
dog.listen_joint([&](const orix::JointData& j) -> bool {
    std::cout << "Joint " << j.id
              << ": pos=" << j.position
              << " vel=" << j.velocity
              << " trq=" << j.torque << "\\n";
    return ++count < 500;
});\
""")

    set_heading(doc, "8.4 电机状态", 2)
    code_block(doc, """\
auto motors = dog.get_motor_status();
for (const auto& m : motors) {
    std::string st = m.online ? "在线" : "离线";
    std::cout << "Joint " << m.id << ": " << st
              << "  temp=" << m.temperature << "C"
              << "  V=" << m.voltage
              << "  errors=" << m.errors.size() << "\\n";
}\
""")

    set_heading(doc, "8.5 电压监控", 2)
    code_block(doc, """\
auto voltages = dog.get_voltage();
double min_v = *std::min_element(voltages.begin(), voltages.end());
double max_v = *std::max_element(voltages.begin(), voltages.end());
std::cout << "电压范围: " << min_v << "V ~ " << max_v << "V\\n";

if (min_v < 18.0) {
    std::cerr << "WARNING: 低压保护即将触发!\\n";
}\
""")

    set_heading(doc, "8.6 综合示例：传感器数据记录到 CSV", 2)
    code_block(doc, """\
#include "orix/orix.h"
#include <atomic>
#include <chrono>
#include <fstream>
#include <iostream>
#include <thread>

int main() {
    orix::OrixClient dog("192.168.66.190");
    std::atomic<bool> running{true};

    std::thread imu_thread([&] {
        std::ofstream csv("imu_log.csv");
        csv << "timestamp_ms,gx,gy,gz,qw,qx,qy,qz\\n";
        dog.listen_imu([&](const orix::ImuData& d) -> bool {
            csv << d.timestamp_ms << ","
                << d.gyroscope.x << "," << d.gyroscope.y
                << "," << d.gyroscope.z << ","
                << d.quaternion.w << "," << d.quaternion.x
                << "," << d.quaternion.y << ","
                << d.quaternion.z << "\\n";
            return running.load();
        });
        std::cout << "IMU 记录完成 -> imu_log.csv\\n";
    });

    std::thread joint_thread([&] {
        std::ofstream csv("joint_log.csv");
        csv << "joint_id,position,velocity,torque,status\\n";
        dog.listen_joint([&](const orix::JointData& j) -> bool {
            csv << j.id << "," << j.position << ","
                << j.velocity << "," << j.torque << ","
                << j.status << "\\n";
            return running.load();
        });
        std::cout << "关节记录完成 -> joint_log.csv\\n";
    });

    std::cout << "开始记录 30 秒...\\n";
    std::this_thread::sleep_for(std::chrono::seconds(30));

    running = false;
    imu_thread.join();
    joint_thread.join();
    std::cout << "记录完成。\\n";
    return 0;
}\
""")
    doc.add_paragraph()

    # ================================================================
    #  第 9 章：API 参考
    # ================================================================
    set_heading(doc, "第九章 API 参考", 1)
    doc.add_paragraph("本章提供 ORIX C++ SDK（orix）全部公开接口的完整参考。")

    # --- 9.1 构造与连接 ---
    set_heading(doc, "9.1 OrixClient 构造与连接", 2)

    set_heading(doc, "9.1.1 构造函数", 3)
    code_block(doc, """\
explicit OrixClient(
    std::string host    = "192.168.66.190",
    int         port    = 13145,
    double      timeout = 5.0);\
""")
    doc.add_paragraph(
        "创建 OrixClient 实例并建立 gRPC channel（惰性连接，首次 RPC 时才真正握手）。"
        "OrixClient 不可拷贝，可移动。"
    )
    add_table(doc, ["参数", "类型", "默认值", "说明"], [
        ["host",    "std::string", '"192.168.66.190"', "机器人 IP 地址"],
        ["port",    "int",         "13145",            "gRPC 服务端口"],
        ["timeout", "double",      "5.0",              "默认 RPC 超时秒数"],
    ])
    code_block(doc, """\
// 默认连接
orix::OrixClient dog("192.168.66.190");

// 自定义参数
orix::OrixClient dog("10.0.0.100", 13145, 10.0);

// unique_ptr 管理
auto dog = std::make_unique<orix::OrixClient>("192.168.66.190");\
""")

    set_heading(doc, "9.1.2 析构函数", 3)
    code_block(doc, "~OrixClient();")
    doc.add_paragraph("析构时自动释放 gRPC channel 和所有相关资源（RAII）。无需手动调用 close()。")

    set_heading(doc, "9.1.3 ping()", 3)
    code_block(doc, "bool ping();")
    doc.add_paragraph("检测机器人是否可达。超时 2 秒，不抛异常。")
    add_table(doc, ["项目", "说明"], [
        ["返回值", "true = 连接正常；false = 不可达或超时"],
        ["超时",   "固定 2 秒，不受构造时 timeout 参数影响"],
        ["异常",   "不抛出任何异常"],
    ])
    code_block(doc, """\
if (!dog.ping()) {
    std::cerr << "无法连接到机器人，请检查网络和服务状态。\\n";
    return 1;
}\
""")

    # --- 9.2 电机控制 ---
    set_heading(doc, "9.2 电机控制", 2)

    set_heading(doc, "9.2.1 enable()", 3)
    code_block(doc, "void enable();")
    doc.add_paragraph("使能全部 16 个电机。调用后电机通电并进入闭环位置控制。")
    add_table(doc, ["项目", "说明"], [
        ["前置条件", "gRPC 连接正常，控制权未被遥控器占用"],
        ["异常",     "ConnectionError --- 超时或连接失败"],
    ])
    code_block(doc, """\
dog.enable();
// 此时状态仍为 Grounded，电机通电但保持坐姿
// 需要调用 stand_up() 才能站起\
""")

    set_heading(doc, "9.2.2 disable()", 3)
    code_block(doc, "void disable();")
    doc.add_paragraph("禁用全部电机（安全停止）。电机立即断电，关节自由落体。")
    add_table(doc, ["项目", "说明"], [
        ["前置条件", "无（任何状态下均可调用）"],
        ["建议",     "在 Grounded 状态下调用，避免机器人从站立姿态跌落"],
        ["异常",     "ConnectionError --- 超时或连接失败"],
    ])
    code_block(doc, """\
// 安全关机序列
dog.walk();        // 停步
dog.sit_down();    // 坐下
// 等待 Grounded ...
dog.disable();     // 禁用电机\
""")

    # --- 9.3 运动指令 ---
    set_heading(doc, "9.3 运动指令", 2)

    set_heading(doc, "9.3.1 walk()", 3)
    code_block(doc, "void walk(double vx = 0.0, double vy = 0.0, double vyaw = 0.0);")
    doc.add_paragraph("发送行走速度指令。")
    add_table(doc, ["参数", "类型", "默认值", "范围", "说明"], [
        ["vx",   "double", "0.0", "[-1, 1]", "前后速度。正值 = 前进，负值 = 后退"],
        ["vy",   "double", "0.0", "[-1, 1]", "左右速度。正值 = 向左，负值 = 向右"],
        ["vyaw", "double", "0.0", "[-1, 1]", "旋转速度。正值 = 逆时针，负值 = 顺时针"],
    ])
    add_table(doc, ["项目", "说明"], [
        ["前置条件", "机器人处于 Standing 或 Walking 状态"],
        ["停步",     "调用 dog.walk() 无参数，或发送 (0, 0, 0)"],
        ["异常",     "ConnectionError, TimeoutError, InvalidStateError"],
    ])
    doc.add_paragraph("速度参数为归一化值，不是物理速度。vx=1.0 表示当前策略允许的最大前进速度。")
    code_block(doc, """\
dog.walk(0.5);                 // 前进
dog.walk(0.0, 0.3);           // 向左平移
dog.walk(0.0, 0.0, 0.3);     // 原地逆时针旋转
dog.walk(0.3, 0.0, 0.2);     // 前进 + 左转
dog.walk();                    // 原地停步\
""")

    set_heading(doc, "9.3.2 stand_up()", 3)
    code_block(doc, "void stand_up();")
    doc.add_paragraph("从坐姿过渡到站立姿态。")
    add_table(doc, ["项目", "说明"], [
        ["前置条件", "机器人处于 Grounded 状态，电机已使能"],
        ["状态变化", "Grounded -> Transitioning -> Standing"],
        ["耗时",     "约 3 秒（150 个控制周期 x 20ms）"],
        ["异常",     "ConnectionError, InvalidStateError"],
    ])
    code_block(doc, """\
dog.enable();
dog.stand_up();
std::this_thread::sleep_for(std::chrono::seconds(4));
// state == RobotState::Standing\
""")

    set_heading(doc, "9.3.3 sit_down()", 3)
    code_block(doc, "void sit_down();")
    doc.add_paragraph("从站立过渡到坐姿。如果正在行走，会先停步再坐下。")
    add_table(doc, ["项目", "说明"], [
        ["前置条件", "机器人处于 Standing 或 Walking 状态"],
        ["状态变化", "Standing/Walking -> Transitioning -> Grounded"],
        ["耗时",     "约 3 秒；从 Walking 坐下约 6 秒（先站稳再坐）"],
        ["异常",     "ConnectionError, InvalidStateError"],
    ])

    set_heading(doc, "9.3.4 set_high_speed() / get_speed_mode()", 3)
    code_block(doc, """\
void        set_high_speed(bool enabled = true);
std::string get_speed_mode();\
""")
    doc.add_paragraph("开启/关闭高速模式（最大 2.5 m/s）。仅在开阔平坦场地使用。")
    add_table(doc, ["项目", "说明"], [
        ["前置条件", "机器人处于 Standing 状态"],
        ["返回值",   'get_speed_mode() 返回 "normal" 或 "high_speed"'],
        ["异常",     "ConnectionError, InvalidStateError"],
    ])
    code_block(doc, """\
dog.set_high_speed(true);        // 启用高速模式
dog.walk(1.0);                    // 全速前进
auto mode = dog.get_speed_mode(); // "high_speed"
dog.set_high_speed(false);        // 恢复正常速度\
""")

    # --- 9.4 状态查询 ---
    set_heading(doc, "9.4 状态查询", 2)

    set_heading(doc, "9.4.1 get_state()", 3)
    code_block(doc, "RobotState get_state();")
    doc.add_paragraph("查询当前机器人 FSM 状态。")
    add_table(doc, ["返回值", "说明"], [
        ["RobotState::Zero",          "初始化/未使能"],
        ["RobotState::Grounded",      "坐姿（安全状态）"],
        ["RobotState::Standing",      "站立，可接收 walk() 或 sit_down()"],
        ["RobotState::Walking",       "行走中，RL 策略以 50Hz 推理驱动步态"],
        ["RobotState::Transitioning", "过渡中，拒绝新的运动指令"],
    ])
    code_block(doc, """\
auto state = dog.get_state();
std::cout << "当前状态: " << orix::to_string(state) << "\\n";
if (state == orix::RobotState::Standing) {
    dog.walk(0.3);
}\
""")

    set_heading(doc, "9.4.2 get_voltage()", 3)
    code_block(doc, "std::vector<double> get_voltage();")
    doc.add_paragraph("读取全部 16 个电机的母线电压（单位 V）。")
    add_table(doc, ["项目", "说明"], [
        ["返回值",   "长度 16 的 double 向量，按关节编号 0-15 排列"],
        ["低压阈值", "18V，低于此值系统自动触发保护坐下"],
    ])
    code_block(doc, """\
auto voltages = dog.get_voltage();
double min_v = *std::min_element(voltages.begin(), voltages.end());
for (int i = 0; i < (int)voltages.size(); ++i) {
    std::string warn = voltages[i] < 18.0 ? " [低压!]" : "";
    std::cout << "  Joint " << i << ": " << voltages[i] << "V" << warn << "\\n";
}\
""")

    set_heading(doc, "9.4.3 get_motor_status()", 3)
    code_block(doc, "std::vector<MotorInfo> get_motor_status();")
    doc.add_paragraph("获取全部 16 个电机的健康状态。")
    code_block(doc, """\
auto motors = dog.get_motor_status();
int online = 0, faults = 0;
for (const auto& m : motors) {
    if (m.online) ++online;
    if (!m.errors.empty()) ++faults;
}
std::cout << "在线: " << online << "/16  故障: " << faults << "/16\\n";\
""")
    doc.add_paragraph("故障码解读：")
    add_table(doc, ["场景", "表现", "处理方式"], [
        ["online=false",      "电机未响应 CAN 通信",  "检查 CAN 线缆和电源"],
        ["errors 非空",       "电机报告故障",          "调用 clear_motor_fault() 清除"],
        ["temperature > 80",  "电机过热",              "降低负载或暂停运动，等待冷却"],
        ["voltage < 18",      "母线电压过低",          "立即充电，系统将触发低压保护"],
    ])

    # --- 9.5 策略管理 ---
    set_heading(doc, "9.5 策略管理", 2)

    set_heading(doc, "9.5.1 get_profile()", 3)
    code_block(doc, "ProfileInfo get_profile();")
    doc.add_paragraph("获取当前运动策略信息及可用策略列表。")
    code_block(doc, """\
auto profile = dog.get_profile();
std::cout << "当前策略: " << profile.current << "\\n";
for (size_t i = 0; i < profile.available.size(); ++i) {
    std::cout << profile.available[i];
    if (i < profile.descriptions.size())
        std::cout << "(" << profile.descriptions[i] << ")";
    std::cout << " ";
}\
""")

    set_heading(doc, "9.5.2 switch_profile()", 3)
    code_block(doc, "ProfileInfo switch_profile(const std::string& name);")
    doc.add_paragraph("切换运动策略。")
    add_table(doc, ["项目", "说明"], [
        ["参数 name", "目标策略名称，必须在 get_profile().available 列表中"],
        ["前置条件",   "机器人处于 Grounded 状态"],
        ["返回值",     "切换后的 ProfileInfo（current 为新策略名）"],
        ["异常",       "ConnectionError, InvalidStateError --- 策略不存在或状态不对"],
    ])

    # --- 9.6 动作系统 ---
    set_heading(doc, "9.6 动作系统", 2)

    set_heading(doc, "9.6.1 list_gestures()", 3)
    code_block(doc, "std::vector<GestureInfo> list_gestures();")
    doc.add_paragraph("列出所有可用预设动作。")
    code_block(doc, """\
auto gestures = dog.list_gestures();
for (const auto& g : gestures) {
    std::cout << "  " << g.name << ": " << g.description
              << " (" << g.duration_ms << "ms)\\n";
}\
""")

    set_heading(doc, "9.6.2 play_gesture()", 3)
    code_block(doc, 'void play_gesture(const std::string& name, double timeout_s = 30.0);')
    doc.add_paragraph("播放预设动作。执行期间机器人处于 Transitioning 状态，不接受新运动指令。")
    add_table(doc, ["项目", "说明"], [
        ["前置条件",      "机器人处于 Standing 状态"],
        ["参数 name",     "动作名称（见 list_gestures()）"],
        ["参数 timeout_s","最长等待时间（秒），默认 30s"],
        ["异常",          "TimeoutError, InvalidStateError"],
    ])
    code_block(doc, """\
dog.stand_up();
std::this_thread::sleep_for(std::chrono::seconds(4));
dog.play_gesture("bow", 7.0);
dog.play_gesture("dance", 35.0);\
""")

    # --- 9.7 诊断与标定 ---
    set_heading(doc, "9.7 诊断与标定", 2)

    set_heading(doc, "9.7.1 clear_motor_fault()", 3)
    code_block(doc, "void clear_motor_fault(const std::vector<int>& joint_ids = {});")
    doc.add_paragraph("清除电机故障码。")
    add_table(doc, ["项目", "说明"], [
        ["参数 joint_ids", "要清除的关节编号列表（0-15）；空向量 = 清除全部"],
        ["前置条件",        "无"],
    ])
    code_block(doc, """\
dog.clear_motor_fault();           // 清除所有
dog.clear_motor_fault({0, 1, 2});  // 只清除前右腿
dog.clear_motor_fault({5});        // 只清除单个关节\
""")

    set_heading(doc, "9.7.2 set_zero()", 3)
    code_block(doc, "void set_zero();")
    doc.add_paragraph("将当前关节位置标定为零点并保存。")
    add_table(doc, ["项目", "说明"], [
        ["前置条件", "机器人处于 Grounded 状态"],
        ["用途",     "出厂标定或维修后重新校准零位"],
        ["注意",     "错误的零位会导致运动异常，非必要不调用"],
    ])

    # --- 9.8 流式接口 ---
    set_heading(doc, "9.8 流式接口", 2)
    doc.add_paragraph(
        "所有 listen_* 方法接受一个回调函数（std::function<bool(...)>），"
        "回调返回 true 继续接收，返回 false 停止流。这些方法是阻塞调用，应在独立的 std::thread 中运行。"
    )

    set_heading(doc, "9.8.1 listen_imu()", 3)
    code_block(doc, "void listen_imu(const std::function<bool(const ImuData&)>& callback);")
    doc.add_paragraph("订阅实时 IMU 数据流（~50Hz）。")
    add_table(doc, ["项目", "说明"], [
        ["产出", "ImuData（含 gyroscope, quaternion, timestamp_ms）"],
        ["频率", "~50 Hz"],
        ["阻塞", "是（在独立线程调用；callback 返回 false 时退出）"],
    ])

    set_heading(doc, "9.8.2 listen_joint()", 3)
    code_block(doc, "void listen_joint(const std::function<bool(const JointData&)>& callback);")
    doc.add_paragraph("订阅实时关节数据流（逐关节上报，~50Hz）。")
    add_table(doc, ["项目", "说明"], [
        ["产出", "JointData（含 id, position, velocity, torque, status）"],
        ["频率", "~50 Hz（逐关节上报，每帧一个关节）"],
        ["阻塞", "是"],
    ])
    code_block(doc, """\
// 要获取完整的 16 关节快照，需自行缓存最近一帧
std::array<orix::JointData, 16> latest{};
dog.listen_joint([&](const orix::JointData& j) -> bool {
    if (j.id >= 0 && j.id < 16)
        latest[j.id] = j;
    return true;
});\
""")

    set_heading(doc, "9.8.3 listen_state()", 3)
    code_block(doc, "void listen_state(const std::function<bool(RobotState)>& callback);")
    doc.add_paragraph("订阅 FSM 状态变更事件（仅在状态切换时触发）。")
    add_table(doc, ["项目", "说明"], [
        ["产出", "RobotState 枚举值"],
        ["频率", "事件驱动（仅状态变化时触发，不定频）"],
        ["阻塞", "是"],
    ])

    # --- 9.9 OrixCamera ---
    set_heading(doc, "9.9 OrixCamera API", 2)
    doc.add_paragraph("OrixCamera 是可选的摄像头模块（依赖 OpenCV >= 4.5），需要链接 orix_camera 库。")
    code_block(doc, '#include "orix/orix_camera.h"')

    set_heading(doc, "9.9.1 构造方式", 3)
    code_block(doc, """\
// 本地 USB 摄像头
explicit OrixCamera(int source=0, int width=640, int height=480, int fps=30);
// 远程 RTSP / MJPEG URL
explicit OrixCamera(const std::string& url);\
""")

    set_heading(doc, "9.9.2 方法速查", 3)
    add_table(doc, ["方法", "返回值", "说明"], [
        ["open()",                              "bool",          "打开摄像头；失败抛 CameraOpenError"],
        ["close()",                             "void",          "释放摄像头及所有资源"],
        ["is_opened()",                         "bool",          "摄像头是否已打开"],
        ["read()",                              "cv::Mat",       "读取一帧（BGR），失败返回空 Mat"],
        ["get_frame_size()",                    "pair<int,int>", "当前画面尺寸 (width, height)"],
        ["save_photo(path)",                    "bool",          "拍照并保存到文件"],
        ["start_recording(path, duration, codec)","void",        "启动后台录像"],
        ["stop_recording()",                    "void",          "停止录像"],
        ["is_recording()",                      "bool",          "是否正在录制"],
        ["stream_mjpeg(port, host)",            "void",          "启动 MJPEG HTTP 流服务器"],
        ["stop_stream()",                       "void",          "停止流服务器"],
        ["is_streaming()",                      "bool",          "流服务器是否运行"],
    ], header_color="D1FAE5")

    set_heading(doc, "9.9.3 MJPEG 流端点", 3)
    add_table(doc, ["URL 路径", "说明"], [
        ["/stream",   "multipart/x-mixed-replace MJPEG 实时流，浏览器可直接查看"],
        ["/snapshot", "单帧 JPEG 快照"],
    ], header_color="D1FAE5")

    # --- 9.10 数据类型参考 ---
    set_heading(doc, "9.10 数据类型参考", 2)

    set_heading(doc, "9.10.1 Vec3", 3)
    doc.add_paragraph("三维向量，用于角速度（rad/s）、速度（m/s）等。")
    code_block(doc, """\
struct Vec3 {
    double x = 0.0;
    double y = 0.0;
    double z = 0.0;
};\
""")

    set_heading(doc, "9.10.2 Quat", 3)
    doc.add_paragraph("四元数（Hamilton 约定：w, x, y, z）。")
    code_block(doc, """\
struct Quat {
    double w = 1.0;    // 实部
    double x = 0.0;    // 虚部 i
    double y = 0.0;    // 虚部 j
    double z = 0.0;    // 虚部 k
};\
""")

    set_heading(doc, "9.10.3 ImuData", 3)
    code_block(doc, """\
struct ImuData {
    Vec3   gyroscope;          // 角速度 (rad/s)
    Quat   quaternion;         // 姿态四元数（Hamilton）
    double timestamp_ms = 0.0; // 时间戳（ms）
};\
""")

    set_heading(doc, "9.10.4 JointData", 3)
    code_block(doc, """\
struct JointData {
    int    id       = 0;    // 关节编号 (0-15)
    double position = 0.0;  // 角位置 (rad)
    double velocity = 0.0;  // 角速度 (rad/s)
    double torque   = 0.0;  // 输出力矩 (N*m)
    int    status   = 0;    // 状态码 (0 = 正常)
};\
""")

    set_heading(doc, "9.10.5 MotorInfo", 3)
    code_block(doc, """\
struct MotorInfo {
    int    id          = 0;
    bool   online      = false;
    double temperature = 0.0;   // 温度 (C)
    double voltage     = 0.0;   // 母线电压 (V)
    double position    = 0.0;   // 编码器位置 (rad)
    double velocity    = 0.0;   // 转速 (rad/s)
    double torque      = 0.0;   // 力矩 (N*m)
    std::vector<int> errors;    // 故障码列表
};\
""")

    set_heading(doc, "9.10.6 ProfileInfo", 3)
    code_block(doc, """\
struct ProfileInfo {
    std::string              current;      // 当前策略名称
    std::vector<std::string> available;    // 可用策略名称
    std::vector<std::string> descriptions; // 策略说明
};\
""")

    set_heading(doc, "9.10.7 GestureInfo", 3)
    code_block(doc, """\
struct GestureInfo {
    std::string name;         // 动作名称
    std::string description;  // 动作描述
    int         duration_ms;  // 预计执行时长（ms）
};\
""")

    set_heading(doc, "9.10.8 RobotState", 3)
    code_block(doc, """\
enum class RobotState {
    Zero          = 0,   // "zero"
    Grounded      = 1,   // "grounded"
    Standing      = 2,   // "standing"
    Walking       = 3,   // "walking"
    Transitioning = 4,   // "transitioning"
    Unknown       = -1,  // "unknown"
};

std::string to_string(RobotState s);\
""")

    # --- 9.11 异常类型 ---
    set_heading(doc, "9.11 异常类型", 2)
    doc.add_paragraph("所有 SDK 异常继承自 orix::OrixError，后者继承自 std::runtime_error。")
    code_block(doc, """\
class OrixError        : public std::runtime_error { ... };
class ConnectionError  : public OrixError { ... };  // gRPC UNAVAILABLE
class InvalidStateError: public OrixError { ... };  // gRPC FAILED_PRECONDITION
class TimeoutError     : public OrixError { ... };  // gRPC DEADLINE_EXCEEDED\
""")
    add_table(doc, ["异常类", "gRPC 状态码", "含义"], [
        ["ConnectionError",   "UNAVAILABLE",        "无法连接机器人（IP 错误、服务未启动）"],
        ["InvalidStateError", "FAILED_PRECONDITION", "机器人状态不满足前提"],
        ["TimeoutError",      "DEADLINE_EXCEEDED",   "RPC 超时（网络延迟、服务无响应）"],
        ["OrixError",         "其他",                "通用 SDK 错误"],
    ])

    # --- 9.12 错误处理 ---
    set_heading(doc, "9.12 错误处理", 2)
    code_block(doc, """\
try {
    orix::OrixClient dog("192.168.66.190");
    dog.stand_up();
    dog.walk(0.5);
} catch (const orix::ConnectionError& e) {
    std::cerr << "连接错误: " << e.what() << "\\n";
} catch (const orix::InvalidStateError& e) {
    std::cerr << "状态错误: " << e.what() << "\\n";
} catch (const orix::TimeoutError& e) {
    std::cerr << "超时: " << e.what() << "\\n";
} catch (const orix::OrixError& e) {
    std::cerr << "SDK 错误: " << e.what() << "\\n";
}\
""")
    doc.add_paragraph("常见 gRPC 错误码：")
    add_table(doc, ["gRPC 状态码", "含义", "常见原因"], [
        ["UNAVAILABLE",        "服务不可达",     "机器人未开机、IP 错误、服务未启动"],
        ["DEADLINE_EXCEEDED",  "请求超时",       "网络延迟、服务端无响应"],
        ["FAILED_PRECONDITION","前置条件不满足", "状态不对（如 Zero 状态下调 walk）"],
        ["PERMISSION_DENIED",  "权限被拒",       "Arbiter 控制权被遥控器占用"],
        ["INTERNAL",           "服务端内部错误", "硬件异常或程序错误"],
    ])

    doc.add_paragraph("Arbiter 控制权冲突处理示例：")
    code_block(doc, """\
auto try_walk = [&](double vx, int retries = 3) {
    for (int i = 0; i < retries; ++i) {
        try {
            dog.walk(vx);
            return true;
        } catch (const orix::OrixError& e) {
            std::cerr << "控制权被遥控器占用，等待 ("
                      << (i+1) << "/" << retries << ")...\\n";
            std::this_thread::sleep_for(std::chrono::seconds(1));
        }
    }
    std::cerr << "无法获得控制权，请确认遥控器已松开摇杆。\\n";
    return false;
};\
""")

    # 附录：方法速查表
    set_heading(doc, "附录：方法速查表", 2)
    add_table(doc, ["分类", "方法", "返回值", "说明"], [
        ["连接", "OrixClient(host, port, timeout)", "--",              "构造并连接"],
        ["",     "~OrixClient()",                    "--",              "析构释放（RAII）"],
        ["",     "ping()",                           "bool",            "连接检测，不抛异常"],
        ["电机", "enable()",                         "void",            "使能全部电机"],
        ["",     "disable()",                        "void",            "禁用全部电机"],
        ["运动", "walk(vx, vy, vyaw)",               "void",            "行走指令，[-1,1]"],
        ["",     "stand_up()",                       "void",            "坐姿 -> 站立"],
        ["",     "sit_down()",                       "void",            "站立 -> 坐姿"],
        ["速度", "set_high_speed(enabled)",           "void",            "高速模式"],
        ["",     "get_speed_mode()",                 "string",          '"normal" 或 "high_speed"'],
        ["状态", "get_state()",                      "RobotState",      "FSM 状态枚举"],
        ["",     "get_voltage()",                    "vector<double>",  "16 路母线电压"],
        ["",     "get_motor_status()",               "vector<MotorInfo>","16 个电机状态"],
        ["策略", "get_profile()",                    "ProfileInfo",     "当前策略及可用列表"],
        ["",     "switch_profile(name)",             "ProfileInfo",     "切换策略（须 Grounded）"],
        ["动作", "list_gestures()",                  "vector<GestureInfo>","可用预设动作列表"],
        ["",     "play_gesture(name, timeout)",      "void",            "播放预设动作"],
        ["诊断", "clear_motor_fault(ids)",           "void",            "清除故障码"],
        ["",     "set_zero()",                       "void",            "标定零位"],
        ["数据流","listen_imu(cb)",                  "void(阻塞)",      "IMU 实时流 ~50Hz"],
        ["",     "listen_joint(cb)",                 "void(阻塞)",      "关节实时流"],
        ["",     "listen_state(cb)",                 "void(阻塞)",      "FSM 状态变更事件"],
    ])
    doc.add_paragraph()

    # ================================================================
    #  第 10 章：安全指南
    # ================================================================
    set_heading(doc, "第十章 安全指南", 1)

    set_heading(doc, "10.1 操作前检查清单", 2)
    doc.add_paragraph("每次启动机器人前，必须逐项确认以下事项：")
    add_table(doc, ["序号", "检查项", "确认方法", "通过标准"], [
        ["1",  "电池电量充足",     "dog.get_voltage()",      "所有电机电压 >= 24V"],
        ["2",  "操作区域净空",     "目视检查",               "周围 1.5 米内无人/易碎物品"],
        ["3",  "机器人放置在地面", "目视检查",               "四脚均接触平坦地面"],
        ["4",  "急停装置就绪",     "确认遥控器在手",         "遥控器已开机"],
        ["5",  "CAN 总线连接牢固", "目视检查",               "CAN 线缆无松动"],
        ["6",  "控制服务已启动",   "SSH 检查",               "brainstem active"],
        ["7",  "电机状态正常",     "get_motor_status()",     "16 电机在线无故障"],
        ["8",  "关节位置合理",     "get_motor_status()",     "绝对值 < 1.0 rad"],
        ["9",  "遥控器优先级确认", "操控遥控器摇杆",         "能正常控制"],
        ["10", "运行环境温度适宜", "环境检查",               "0-40 C"],
    ])
    doc.add_paragraph()
    doc.add_paragraph("启动前检查代码：")
    code_block(doc, """\
#include "orix/orix.h"
#include <algorithm>
#include <cmath>
#include <iostream>

int main() {
    orix::OrixClient dog("192.168.66.190");

    // 1. 检查电压
    auto voltages = dog.get_voltage();
    double min_v = *std::min_element(voltages.begin(), voltages.end());
    std::cout << "电压: " << min_v << "V -- "
              << (min_v >= 24.0 ? "通过" : "不通过!") << "\\n";

    // 2. 检查电机状态
    auto motors = dog.get_motor_status();
    int offline = 0, faulted = 0;
    for (const auto& m : motors) {
        if (!m.online) ++offline;
        if (!m.errors.empty()) ++faulted;
    }
    std::cout << "电机在线: " << (16 - offline) << "/16 -- "
              << (offline == 0 ? "通过" : "不通过!") << "\\n";
    std::cout << "电机故障: " << faulted << "/16 -- "
              << (faulted == 0 ? "通过" : "需清除") << "\\n";

    // 3. 检查关节位置
    bool pos_ok = true;
    for (const auto& m : motors) {
        if (std::abs(m.position) > 1.0) {
            std::cout << "关节 " << m.id << " 位置异常: "
                      << m.position << " rad\\n";
            pos_ok = false;
        }
    }
    std::cout << "关节位置: " << (pos_ok ? "通过" : "异常!") << "\\n";

    // 4. 检查状态
    auto state = dog.get_state();
    std::cout << "当前状态: " << orix::to_string(state) << "\\n";
    return 0;
}\
""")

    set_heading(doc, "10.2 低压保护", 2)
    doc.add_paragraph("当任一电机报告的母线电压低于 18V 时，系统自动触发低压保护：")
    numbered(doc, "检测到任一电机母线电压低于 18V")
    numbered(doc, "自动发送 SitDown 指令")
    numbered(doc, "等待机器人进入 Grounded 状态")
    numbered(doc, "禁用全部电机")
    doc.add_paragraph()
    add_table(doc, ["电压范围", "状态", "建议"], [
        [">= 24V",    "正常",   "正常使用"],
        ["20V - 24V", "偏低",   "关注电量"],
        ["18V - 20V", "低电量", "尽快结束任务"],
        ["< 18V",     "危险",   "自动触发保护"],
    ])

    set_heading(doc, "10.3 关节限位保护", 2)
    doc.add_paragraph(
        "每个关节的位置都受到绝对限位保护。默认限位为 3.14 rad（约 180 度）。"
        "任一关节位置绝对值超过此阈值将立即触发故障（Fault）。"
    )
    add_table(doc, ["配置", "环境变量", "默认值"], [
        ["关节限位", "HAN_DOG_JOINT_LIMIT_RAD", "3.14 rad"],
    ])
    doc.add_paragraph("触发后果：")
    numbered(doc, "对应电机立即进入故障状态")
    numbered(doc, "机器人可能失去该腿的控制，导致摔倒")
    numbered(doc, "必须手动清除故障后才能恢复")
    doc.add_paragraph("预防措施：")
    bullet(doc, "确保零位标定正确（set_zero() 在标准坐姿下执行）")
    bullet(doc, "不要在关节卡住或受阻的情况下强行运动")
    code_block(doc, """\
# 设置更严格的限位（在机器狗上设置）
export HAN_DOG_JOINT_LIMIT_RAD=2.5   # 约 143 度\
""")

    set_heading(doc, "10.4 电机故障处理", 2)

    set_heading(doc, "10.4.1 检测故障", 3)
    code_block(doc, """\
auto motors = dog.get_motor_status();
for (const auto& m : motors) {
    if (!m.errors.empty())
        std::cout << "Joint " << m.id << " 故障码: ";
    for (int e : m.errors) std::cout << e << " ";
    if (!m.online)
        std::cout << "Joint " << m.id << " 离线 (CAN 通信中断)";
}\
""")

    set_heading(doc, "10.4.2 清除故障", 3)
    code_block(doc, """\
// 方法 1：清除所有电机故障
dog.clear_motor_fault();

// 方法 2：只清除特定关节
dog.clear_motor_fault({0, 1, 2});   // 只清除前右腿\
""")

    set_heading(doc, "10.4.3 故障处理流程", 3)
    numbered(doc, "判断机器人是否正在运动中。")
    numbered(doc, "如果在运动中：先调用 dog.sit_down() 安全坐下，等待进入 Grounded 状态，再调用 dog.clear_motor_fault() 清除故障，最后重新检查电机状态。")
    numbered(doc, "如果未在运动中：直接调用 dog.clear_motor_fault() 清除故障，然后重新检查。")
    numbered(doc, "如果故障清除成功，可恢复正常运动；如果故障持续存在，需检查硬件。")

    set_heading(doc, "10.4.4 反复出现的故障", 3)
    doc.add_paragraph("如果 clear_motor_fault() 后故障立即重新出现，说明存在持续性硬件问题：")
    add_table(doc, ["表现", "可能原因", "处理"], [
        ["单个电机反复故障",   "CAN 线缆松动",   "重新插拔对应 CAN 接头"],
        ["多个电机同时故障",   "母线问题",        "检查电池和配电板"],
        ["故障码伴随高温",     "电机过热",        "停止运动，等待冷却至 50C 以下"],
        ["关节位置跳变后故障", "编码器异常",      "重新标零 set_zero()"],
    ])

    set_heading(doc, "10.5 急停程序", 2)

    set_heading(doc, "10.5.1 软件急停（SDK）", 3)
    code_block(doc, """\
// 优雅急停
dog.walk();          // 停步
dog.sit_down();      // 坐下
// 等待 Grounded ...
dog.disable();       // 禁用电机\
""")

    set_heading(doc, "10.5.2 遥控器急停", 3)
    doc.add_paragraph("立即操控遥控器摇杆——遥控器自动获得最高控制权，接管机器人。")
    warning_box(doc, "遥控器始终具有最高优先级，可在任何时刻覆盖 SDK 指令。")

    set_heading(doc, "10.5.3 硬件断电", 3)
    doc.add_paragraph("最紧急时直接关闭电池开关或拔出电池。注意断电跌落可能造成机械损伤。")

    set_heading(doc, "10.5.4 急停优先级", 3)
    add_table(doc, ["优先级", "方式", "说明"], [
        ["1", "硬件断电",   "物理级，不可覆盖"],
        ["2", "遥控器控制", "Arbiter 保证最高软件优先级"],
        ["3", "SDK disable()", "gRPC 调用"],
        ["4", "低压自动保护", "系统自动触发"],
    ])

    set_heading(doc, "10.6 Arbiter 控制仲裁安全", 2)
    doc.add_paragraph("ORIX 的 ControlArbiter（控制仲裁器）实施以下安全规则：")
    add_table(doc, ["规则", "说明"], [
        ["遥控器优先", "遥控器的优先级永远高于 SDK 指令，3 秒无操作自动释放"],
        ["超时释放",   "遥控器停止操作 3 秒后，gRPC 自动获得控制权"],
        ["无抢占",     "gRPC 不能强行抢占遥控器的控制权"],
        ["超时可配",   "超时时间通过 HAN_DOG_ARBITER_TIMEOUT 配置，默认 3 秒"],
    ])
    doc.add_paragraph("对 SDK 开发者的影响：")
    bullet(doc, "当遥控器在线并活跃时，SDK 的运动指令将被忽略")
    bullet(doc, "状态查询（get_state(), get_voltage() 等）不受 Arbiter 影响，始终可用")
    bullet(doc, "遥控器释放后约 3 秒，SDK 自动恢复控制权，无需额外操作")
    doc.add_paragraph("安全启示：")
    bullet(doc, "现场测试时必须携带遥控器，作为最终的安全手段")
    bullet(doc, "不要依赖 SDK 作为唯一的急停方式")
    bullet(doc, "如果 SDK 指令无响应，先检查遥控器是否在线")

    set_heading(doc, "10.7 禁止操作", 2)
    doc.add_paragraph("以下操作严格禁止：")
    add_table(doc, ["序号", "禁止操作", "原因"], [
        ["1", "在桌面/高台上运行",           "摔落损坏"],
        ["2", "电机使能时翻转机器人",         "电机对抗损坏齿轮箱"],
        ["3", "电机使能时插拔 CAN 线",        "CAN 总线短路"],
        ["4", "长时间不监控温度",             "电机过热损坏绕组"],
        ["5", "关节卡死时反复运动",           "烧毁电机"],
        ["6", "多客户端同时发运动指令",       "指令冲突抖动"],
    ])

    set_heading(doc, "10.8 标准关机序列", 2)
    code_block(doc, """\
#include "orix/orix.h"
#include <chrono>
#include <thread>

void safe_shutdown(orix::OrixClient& dog) {
    try {
        dog.walk();
        dog.sit_down();
        for (int i = 0; i < 50; ++i) {
            if (dog.get_state() == orix::RobotState::Grounded)
                break;
            std::this_thread::sleep_for(
                std::chrono::milliseconds(200));
        }
        dog.disable();
    } catch (...) {
        try { dog.disable(); } catch (...) {}
    }
}\
""")

    set_heading(doc, "10.9 故障诊断树", 2)
    doc.add_paragraph("遇到机器人异常时，按以下步骤逐级排查：")
    numbered(doc, "机器人是否有响应？无响应：检查电源、网络连通性、CAN 线缆、SSH 连通、brainstem 服务。有响应：进入下一步。")
    numbered(doc, "机器人是否在运动中？在运动中：检查运动是否正常。未运动：调用 dog.get_state() 查询状态。")
    numbered(doc, "有故障码：调用 dog.clear_motor_fault() 清除后重新检查。")
    numbered(doc, "无故障码：检查状态机当前状态是否正确。")
    doc.add_paragraph("运动异常诊断：")
    numbered(doc, "机器人摔倒：首先用 get_voltage() 检查电压。低于 18V 说明是低压保护触发；电压正常则检查关节限位。")
    numbered(doc, "运动不流畅：用 get_motor_status() 检查电机温度。超过 70C 说明过热。")
    numbered(doc, "指令无效：检查 Arbiter 状态，确认遥控器是否在线占用控制权。")

    set_heading(doc, "10.10 安全相关环境变量", 2)
    add_table(doc, ["环境变量", "默认值", "说明"], [
        ["HAN_DOG_JOINT_LIMIT_RAD",      "3.14",  "关节绝对限位 (rad)"],
        ["HAN_DOG_ARBITER_TIMEOUT",      "3",     "遥控器释放后控制权等待时间 (s)"],
        ["HAN_DOG_SHUTDOWN_TIMEOUT",     "8",     "关机等待超时 (s)"],
        ["HAN_DOG_STARTUP_TIMEOUT",      "10",    "FSM 启动等待 Grounded 超时 (s)"],
        ["HAN_DOG_SENSOR_LOW_THRESHOLD", "3",     "传感器低读数计数阈值"],
    ])

    warning_box(doc,
        "ORIX 安全操作五条铁律：(1) 永远带遥控器 --- 遥控器是最终安全手段；"
        "(2) 先坐再断电 --- disable() 前必须 sit_down()；"
        "(3) 关注电压 --- 低于 24V 准备收工；"
        "(4) 监控温度 --- 超过 70C 立即休息；"
        "(5) 地面运行 --- 永远在地面操作。"
    )
    doc.add_paragraph()

    # ================================================================
    #  第 11 章：常见问题
    # ================================================================
    set_heading(doc, "第十一章 常见问题", 1)

    # --- 11.1 连接问题 ---
    set_heading(doc, "11.1 连接问题", 2)

    set_heading(doc, "11.1.1 Q: 连接超时 (TimeoutError)", 3)
    doc.add_paragraph("排查步骤：")
    code_block(doc, """\
# 1. 确认网络可达
ping 192.168.66.190
# 2. 确认端口开放
nc -zv 192.168.66.190 13145
# 3. 确认服务运行
ssh sunrise@192.168.66.190 "systemctl status brainstem"\
""")
    add_table(doc, ["原因", "解决"], [
        ["机器人未开机",   "打开电池开关，等 30 秒"],
        ["不在同一网段",   "确认局域网连通"],
        ["brainstem 未启动","systemctl start brainstem"],
        ["防火墙拦截",     "放行 13145 端口"],
        ["超时太短",       "构造时增大 timeout"],
    ])

    set_heading(doc, "11.1.2 Q: 连接被拒绝 (ConnectionError)", 3)
    code_block(doc, """\
try {
    orix::OrixClient dog("192.168.66.190");
    dog.get_state();
} catch (const orix::ConnectionError& e) {
    std::cerr << "连接失败: " << e.what() << "\\n";
    // 检查 brainstem 服务、IP 地址和端口
}\
""")

    set_heading(doc, "11.1.3 Q: IP 地址不确定", 3)
    code_block(doc, """\
# 方法 1：通过路由器管理页面查看 DHCP 租约
# 方法 2：在已知网段扫描 gRPC 端口
nmap -p 13145 192.168.66.0/24
# 方法 3：SSH 登录后查看 IP
ip addr show\
""")
    doc.add_paragraph("默认 IP 为 192.168.66.190，如果网络配置未更改，直接使用即可。")

    # --- 11.2 编译问题 ---
    set_heading(doc, "11.2 编译问题", 2)

    set_heading(doc, "11.2.1 Q: CMake 找不到 gRPC", 3)
    code_block(doc, """\
cmake .. -DCMAKE_PREFIX_PATH=/path/to/grpc/install
# 或使用 vcpkg
cmake .. -DCMAKE_TOOLCHAIN_FILE=/path/to/vcpkg/scripts/buildsystems/vcpkg.cmake\
""")

    set_heading(doc, "11.2.2 Q: 链接错误 undefined reference to grpc", 3)
    code_block(doc, """\
target_link_libraries(my_app PRIVATE
    orix::orix
    gRPC::grpc++
    protobuf::libprotobuf
    pthread
)\
""")

    set_heading(doc, "11.2.3 Q: C++17 标准不满足", 3)
    code_block(doc, """\
g++ --version  # 需要 GCC 9+
# CMakeLists.txt:
set(CMAKE_CXX_STANDARD 17)
set(CMAKE_CXX_STANDARD_REQUIRED ON)\
""")

    # --- 11.3 电机不动 ---
    set_heading(doc, "11.3 电机不动", 2)

    set_heading(doc, "11.3.1 Q: enable() 后电机没反应", 3)
    code_block(doc, """\
// 1. 检查状态
auto state = dog.get_state();
std::cout << "状态: " << orix::to_string(state) << "\\n";

// 2. 检查电机在线
auto motors = dog.get_motor_status();
for (const auto& m : motors) {
    if (!m.online)
        std::cout << "Joint " << m.id << " 离线\\n";
    if (!m.errors.empty())
        std::cout << "Joint " << m.id << " 有故障\\n";
}

// 3. 检查电压
auto voltages = dog.get_voltage();
double min_v = *std::min_element(voltages.begin(), voltages.end());
std::cout << "最低电压: " << min_v << "V\\n";\
""")

    set_heading(doc, "11.3.2 Q: 只有部分电机动，其他不动", 3)
    doc.add_paragraph("常见原因：")
    bullet(doc, "对应腿的 CAN 线缆断开——检查该腿的 CAN 总线物理连接")
    bullet(doc, "单个电机驱动板故障——更换后重试")
    bullet(doc, "电机过热保护——查看温度值，等待冷却")

    set_heading(doc, "11.3.3 Q: enable() 后状态没变化", 3)
    doc.add_paragraph("enable() 使能电机，但不会改变 FSM 状态。使能后需要手动切换状态：")
    code_block(doc, """\
dog.enable();
// 此时状态仍为 Grounded，电机通电但保持坐姿
dog.stand_up();
// 现在状态变为 Standing\
""")

    # --- 11.4 机器人摔倒 ---
    set_heading(doc, "11.4 机器人摔倒", 2)

    set_heading(doc, "11.4.1 Q: 行走中突然坐下", 3)
    doc.add_paragraph("最可能原因：低压保护触发。")
    code_block(doc, """\
auto voltages = dog.get_voltage();
double min_v = *std::min_element(voltages.begin(), voltages.end());
std::cout << "当前电压: " << min_v << "V\\n";
// 如果 < 18V，说明是低压保护\
""")
    doc.add_paragraph("其他可能原因：")
    add_table(doc, ["原因", "诊断方法"], [
        ["低电压 (< 18V)",    "get_voltage() 返回值低于 18"],
        ["电机过热",          "get_motor_status() 中 temperature > 80C"],
        ["关节超限",          "get_motor_status() 中 position 绝对值过大"],
        ["CAN 通信中断",      "get_motor_status() 中部分电机 online=false"],
    ])

    set_heading(doc, "11.4.2 Q: 站起来后立即摔倒", 3)
    doc.add_paragraph("排查：")
    numbered(doc, "地面是否太滑——在有摩擦力的地面（橡胶垫、地毯）上测试")
    numbered(doc, "零位是否正确——如果关节零位偏移，站立姿态会畸形")
    numbered(doc, "策略是否匹配——确认加载的运动策略与当前机器人硬件配置一致")

    # --- 11.5 行走指令无效 ---
    set_heading(doc, "11.5 行走指令无效", 2)

    set_heading(doc, "11.5.1 Q: walk() 不生效", 3)
    doc.add_paragraph("必须按状态机顺序操作：")
    code_block(doc, """\
auto state = dog.get_state();
if (state == orix::RobotState::Grounded) {
    std::cout << "需要先站起来\\n";
    dog.stand_up();
    std::this_thread::sleep_for(std::chrono::seconds(4));
}
if (dog.get_state() == orix::RobotState::Standing) {
    dog.walk(0.3);  // 现在应该生效
}\
""")

    set_heading(doc, "11.5.2 Q: 遥控器在线时 SDK 指令被忽略", 3)
    doc.add_paragraph(
        "这是正常行为。ControlArbiter 保证遥控器最高优先级。"
        "关闭遥控器或松开摇杆 3 秒后 SDK 自动恢复控制权。"
    )

    # --- 11.6 策略切换失败 ---
    set_heading(doc, "11.6 策略切换失败", 2)

    set_heading(doc, "11.6.1 Q: switch_profile() 报错", 3)
    doc.add_paragraph("前置条件：策略切换只能在 Grounded 状态下执行。")
    code_block(doc, """\
auto state = dog.get_state();
if (state != orix::RobotState::Grounded) {
    dog.walk();
    dog.sit_down();
    for (int i = 0; i < 50; ++i) {
        if (dog.get_state() == orix::RobotState::Grounded) break;
        std::this_thread::sleep_for(std::chrono::milliseconds(200));
    }
}
auto profile = dog.switch_profile("mini");\
""")

    set_heading(doc, "11.6.2 Q: 策略名称不存在", 3)
    code_block(doc, """\
auto profile = dog.get_profile();
std::cout << "可用策略: ";
for (const auto& p : profile.available) std::cout << p << " ";
std::cout << "\\n";
dog.switch_profile(profile.available[0]);\
""")

    set_heading(doc, "11.6.3 Q: 切换策略后行为异常", 3)
    doc.add_paragraph("每个策略针对特定硬件配置训练。切换后如果出现异常运动：")
    numbered(doc, "确认策略与当前硬件配置匹配")
    numbered(doc, "立即坐下：dog.sit_down()")
    numbered(doc, "切回之前正常工作的策略")

    # --- 11.7 IMU 数据异常 ---
    set_heading(doc, "11.7 IMU 数据异常", 2)

    set_heading(doc, "11.7.1 Q: 四元数值看起来不对", 3)
    doc.add_paragraph("ORIX SDK 使用 Hamilton 约定，四元数顺序为 (w, x, y, z)。如果你的应用使用 ROS 约定 (x, y, z, w)，需要手动转换。")
    code_block(doc, """\
// SDK (Hamilton) -> ROS
// ros_quat = {d.quaternion.x, d.quaternion.y,
//             d.quaternion.z, d.quaternion.w}\
""")

    set_heading(doc, "11.7.2 Q: 机器人静止但角速度不为零", 3)
    doc.add_paragraph(
        "IMU 的陀螺仪存在零偏（bias），静止时读数不会严格为零。"
        "典型零偏范围：|gyro.x| < 0.02 rad/s 属于正常；|gyro.x| > 0.05 rad/s 可能需要重新标定。"
        "如需高精度角速度，可在启动时采集 100 帧静止数据取均值作为零偏。"
    )

    set_heading(doc, "11.7.3 Q: listen_imu() 频率不是 50Hz", 3)
    doc.add_paragraph("IMU 标称频率约 50Hz，实际可能在 48-52Hz 波动，这是正常的。如果频率显著偏低（< 30Hz），检查：")
    bullet(doc, "串口连接是否正常")
    bullet(doc, "brainstem 服务是否负载过高")
    bullet(doc, "网络延迟（gRPC 传输开销）")

    # --- 11.8 流式接口问题 ---
    set_heading(doc, "11.8 流式接口问题", 2)

    set_heading(doc, "11.8.1 Q: listen_imu() 阻塞了主线程", 3)
    doc.add_paragraph("所有 listen_* 方法都是阻塞的。使用独立线程：")
    code_block(doc, """\
std::atomic<bool> running{true};
std::thread imu_thread([&] {
    dog.listen_imu([&](const orix::ImuData& d) -> bool {
        // 处理数据...
        return running.load();
    });
});
// 主线程继续执行其他操作...
running = false;
imu_thread.join();\
""")

    set_heading(doc, "11.8.2 Q: 流中断后如何重连", 3)
    doc.add_paragraph("当网络断开或服务重启时，流式回调会结束。重新创建 OrixClient 并重新订阅：")
    code_block(doc, """\
while (true) {
    try {
        orix::OrixClient dog("192.168.66.190");
        dog.listen_imu([](const orix::ImuData& d) -> bool {
            // 处理数据...
            return true;
        });
    } catch (const orix::ConnectionError& e) {
        std::cerr << "连接断开，5 秒后重试...\\n";
        std::this_thread::sleep_for(std::chrono::seconds(5));
    }
}\
""")

    # --- 11.9 性能与延迟 ---
    set_heading(doc, "11.9 性能与延迟", 2)

    set_heading(doc, "11.9.1 Q: C++ SDK 比 Python 快多少？", 3)
    add_table(doc, ["指标", "Python SDK", "C++ SDK"], [
        ["单次 RPC 延迟",  "1-5 ms",    "0.5-2 ms"],
        ["流式数据处理",   "受 GIL 限制","真正多线程并行"],
        ["启动时间",       "~500 ms",    "~50 ms"],
        ["内存占用",       "~50 MB",     "~5 MB"],
        ["适用场景",       "快速原型、教学","嵌入式部署、高频控制"],
    ])

    set_heading(doc, "11.9.2 Q: 50Hz 控制频率够用吗？", 3)
    doc.add_paragraph(
        "ORIX 的运动控制循环固定为 50Hz。SDK 发送的指令在下一个控制周期生效。"
        "不需要以超过 50Hz 的频率发送 walk() 指令。C++ 的低延迟特性在需要精确时序控制的场景中尤为重要。"
    )

    # --- 11.10 环境变量参考 ---
    set_heading(doc, "11.10 环境变量参考", 2)
    doc.add_paragraph("以下环境变量在机器人端（brainstem 服务）生效，SDK 客户端不直接使用。")
    add_table(doc, ["环境变量", "类型", "默认值", "说明"], [
        ["HAN_DOG_PORT",                "int",   "13145",    "gRPC 服务监听端口"],
        ["HAN_DOG_CONTROLLER",          "str",   "auto",     "遥控器设备（自动检测）"],
        ["HAN_DOG_DEFAULT_PROFILE",     "str",   "（无）",   "默认加载的运动策略"],
        ["HAN_DOG_PROFILE_DIR",         "str",   "profiles/","策略文件目录"],
        ["HAN_DOG_ARBITER_TIMEOUT",     "int",   "3",        "遥控器释放后控制权切换 (s)"],
        ["HAN_DOG_SENSOR_LOW_THRESHOLD","int",   "3",        "传感器低读数计数阈值"],
        ["HAN_DOG_SHUTDOWN_TIMEOUT",    "int",   "8",        "关机等待超时 (s)"],
        ["HAN_DOG_STARTUP_TIMEOUT",     "int",   "10",       "启动等待 Grounded 超时 (s)"],
        ["HAN_DOG_JOINT_LIMIT_RAD",     "float", "3.14",     "关节绝对限位 (rad)"],
        ["HAN_DOG_LOG_DIR",             "str",   "logs",     "日志文件目录"],
        ["HAN_DOG_LOG",                 "str",   "INFO",     "日志级别"],
        ["MEDULLA_STANDALONE",          "str",   "（无）",   "1 = 启用独立模式"],
    ])

    # --- 11.11 摄像头问题 ---
    set_heading(doc, "11.11 摄像头问题", 2)

    set_heading(doc, "11.11.1 Q: OrixCamera 编译失败", 3)
    doc.add_paragraph("OrixCamera 需要 OpenCV >= 4.5。确认安装：")
    code_block(doc, """\
# vcpkg
./vcpkg install opencv4
# 或使用系统包管理器
sudo apt install libopencv-dev\
""")

    set_heading(doc, "11.11.2 Q: 远程访问摄像头", 3)
    doc.add_paragraph("开发机上可通过 RTSP URL 访问机器人摄像头：")
    code_block(doc, """\
orix::OrixCamera cam("rtsp://192.168.66.190:8554/cam");
cam.open();
cv::Mat frame = cam.read();\
""")

    # --- 11.12 网络排障 ---
    set_heading(doc, "11.12 网络排障", 2)
    code_block(doc, """\
# 1. Ping 机器人
ping 192.168.66.190
# 2. 检查端口
nc -zv 192.168.66.190 13145
# 3. 检查服务
ssh sunrise@192.168.66.190 "systemctl status brainstem"
# 4. 远程开发（SSH 隧道）
ssh -L 13145:192.168.66.190:13145 sunrise@192.168.66.190
# 然后连接 localhost\
""")

    # --- 11.13 提交 Bug 报告 ---
    set_heading(doc, "11.13 提交 Bug 报告", 2)
    doc.add_paragraph("遇到无法自行解决的问题时，请收集以下信息：")
    code_block(doc, """\
1. SDK 版本：orix C++ SDK v1.0.0
2. 编译器版本：g++ --version / cl --version
3. 操作系统：(例：Ubuntu 22.04 / Windows 11)
4. 机器人固件版本：(SSH 登录后查看 brainstem 版本)
5. 问题描述：(清晰描述预期行为和实际行为)
6. 复现步骤：(从连接开始，逐步列出操作)
7. 错误信息：(完整的异常信息或错误输出)
8. 机器人日志：(提供 logs/han_dog_YYYYMMDD.log 中相关片段)\
""")

    doc.add_paragraph("收集诊断信息：")
    code_block(doc, """\
#include "orix/orix.h"
#include <iostream>

int main() {
    std::cout << "=== ORIX C++ SDK 诊断信息 ===\\n";
    std::cout << "SDK 版本: 1.0.0\\n";

    try {
        orix::OrixClient dog("192.168.66.190", 13145, 5.0);

        auto state = dog.get_state();
        std::cout << "连接:      成功\\n";
        std::cout << "机器人状态: " << orix::to_string(state) << "\\n";

        auto voltages = dog.get_voltage();
        if (!voltages.empty()) {
            double min_v = *std::min_element(
                voltages.begin(), voltages.end());
            double max_v = *std::max_element(
                voltages.begin(), voltages.end());
            std::cout << "电压范围:  " << min_v << "V ~ "
                      << max_v << "V\\n";
        }

        auto motors = dog.get_motor_status();
        int online = 0, faults = 0;
        double max_temp = 0;
        for (const auto& m : motors) {
            if (m.online) ++online;
            if (!m.errors.empty()) ++faults;
            if (m.temperature > max_temp) max_temp = m.temperature;
        }
        std::cout << "电机在线:  " << online << "/16\\n";
        std::cout << "电机故障:  " << faults << "/16\\n";
        std::cout << "最高温度:  " << max_temp << "C\\n";

        auto profile = dog.get_profile();
        std::cout << "当前策略:  " << profile.current << "\\n";
    } catch (const orix::ConnectionError& e) {
        std::cerr << "连接:      失败 - " << e.what() << "\\n";
    } catch (const std::exception& e) {
        std::cerr << "异常:      " << e.what() << "\\n";
    }

    std::cout << "==============================\\n";
    std::cout << "请将以上输出复制到 Bug 报告中。\\n";
    return 0;
}\
""")

    doc.add_paragraph("日志收集：")
    code_block(doc, """\
# SSH 登录机器人
ssh sunrise@192.168.66.190

# 查看今天的日志
cat logs/han_dog_$(date +%Y%m%d).log

# 只看错误和警告
grep -E "WARNING|SEVERE" logs/han_dog_$(date +%Y%m%d).log

# 导出日志到本地
scp sunrise@192.168.66.190:logs/han_dog_*.log ./han_dog_logs/\
""")

    # --- 11.14 SSH 访问与系统服务 ---
    set_heading(doc, "11.14 SSH 访问与系统服务", 2)
    doc.add_paragraph("默认账户：")
    add_table(doc, ["项目", "默认值"], [
        ["用户名",   "sunrise"],
        ["密码",     "sunrise"],
        ["IP 地址",  "192.168.66.190"],
        ["SSH 端口", "22"],
    ])
    code_block(doc, """\
ssh sunrise@192.168.66.190
# 提示输入密码：sunrise\
""")

    doc.add_paragraph("brainstem 服务管理：")
    code_block(doc, """\
# 查看服务状态
systemctl status brainstem

# 查看最近 100 行日志
journalctl -u brainstem -n 100

# 实时跟踪日志
journalctl -u brainstem -f

# 重启服务
sudo systemctl restart brainstem

# 停止服务
sudo systemctl stop brainstem

# 启动服务
sudo systemctl start brainstem\
""")

    doc.add_paragraph("无网络环境下的备选连接方式：")
    numbered(doc, "用网线将开发机直接连接到机器人的 RJ45 口")
    numbered(doc, "将开发机的以太网接口配置为静态 IP，例如 192.168.66.100/24")
    numbered(doc, "使用默认地址连接：ssh sunrise@192.168.66.190")
    doc.add_paragraph("这种方式不依赖任何路由器或交换机，是调试现场最可靠的连接手段。")

    # --- 11.15 快速排障速查表 ---
    set_heading(doc, "11.15 快速排障速查表", 2)
    add_table(doc, ["症状", "首先检查", "可能原因", "解决方案"], [
        ["连接超时",       "ping 机器人",          "网络不通",     "检查网线/Wi-Fi/IP"],
        ["连接被拒",       "brainstem 服务",       "服务未启动",   "systemctl start brainstem"],
        ["编译失败",       "CMake 配置",           "缺少 gRPC",    "安装 vcpkg 依赖"],
        ["enable() 无效",  "get_motor_status()",   "电机离线",     "检查 CAN 线"],
        ["walk() 无效",    "get_state()",          "不在 Standing","先 stand_up()"],
        ["walk() 被忽略",  "遥控器状态",           "Arbiter 被占用","松开遥控器"],
        ["突然坐下",       "get_voltage()",        "低压 (< 18V)", "充电"],
        ["流阻塞主线程",   "代码结构",             "未用独立线程", "std::thread"],
        ["策略切换失败",   "get_state()",          "不在 Grounded","先 sit_down()"],
    ])
    doc.add_paragraph()

    # ── 尾页 ─────────────────────────────────────────────────
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        "ORIX C++ SDK v1.0.0  |  大算机器人  |  机密文件，请勿外传"
    ).font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    return doc


if __name__ == "__main__":
    doc = build_doc()
    out = "orix_cpp_sdk_manual.docx"
    doc.save(out)
    print(f"已生成：{out}")
